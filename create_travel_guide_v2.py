from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153) if level == 1 else RGBColor(0, 0, 0)
    return heading

def add_paragraph_custom(doc, text, bold=False, size=10.5, color=None):
    """添加自定义段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_bullet_list(doc, items):
    """添加项目符号列表"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.runs[0] if p.runs else p.add_run(item)
        if not p.runs:
            run.text = item
        run.font.name = '微软雅黑'
        run.font.size = Pt(10.5)

def add_numbered_list(doc, items):
    """添加编号列表"""
    for item in items:
        p = doc.add_paragraph(style='List Number')
        run = p.runs[0] if p.runs else p.add_run(item)
        if not p.runs:
            run.text = item
        run.font.name = '微软雅黑'
        run.font.size = Pt(10.5)

def create_travel_guide_v2():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    
    # 封面标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('东南亚 · 日韩\n旅游完全指南')
    run.font.name = '微软雅黑'
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('\n2025-2026 最新实用版\n行程·餐厅·酒店·交通·购物·避坑 一本搞定')
    run.font.name = '微软雅黑'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_page_break()
    
    # 目录
    add_heading_custom(doc, '目 录', level=1)
    toc_items = [
        '一、泰国完全指南',
        '   1.1 曼谷3日详细行程',
        '   1.2 清迈3日详细行程',
        '   1.3 普吉岛3日详细行程',
        '   1.4 餐厅推荐（附地址价格）',
        '   1.5 酒店推荐（分区域分价位）',
        '   1.6 交通价格表',
        '   1.7 避坑清单',
        '   1.8 购物清单',
        '   1.9 实用泰语',
        '二、新加坡完全指南',
        '三、马来西亚完全指南',
        '四、日本完全指南',
        '五、韩国完全指南',
        '六、万能行李清单',
        '七、紧急联系大全'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.3 if item.startswith('   ') else 0)
    
    doc.add_page_break()
    
    # ==================== 泰国 ====================
    add_heading_custom(doc, '一、泰国完全指南', level=1)
    
    # ===== 曼谷3日行程 =====
    add_heading_custom(doc, '1.1 曼谷3日详细行程（可直接照搬）', level=2)
    
    add_paragraph_custom(doc, '【Day 1：经典景点日】', bold=True, size=11)
    add_bullet_list(doc, [
        '08:00 酒店早餐',
        '09:00 大皇宫（门票500泰铢，游玩2小时，注意穿长裤）',
        '11:30 卧佛寺（门票200泰铢，看巨大卧佛）',
        '13:00 午餐：Thip Samai（鬼门炒粉，人均150泰铢）',
        '15:00 郑王庙（门票100泰铢，拍照圣地）',
        '17:00 湄南河游船（公交船15泰铢，观光船150泰铢）',
        '19:00 考山路夜市（晚餐+按摩，人均300泰铢）'
    ])
    
    add_paragraph_custom(doc, '【Day 2：购物美食日】', bold=True, size=11)
    add_bullet_list(doc, [
        '09:00 水门市场（早市，衣服便宜）',
        '11:00 Central World商圈（曼谷最大商场）',
        '13:00 午餐：Kub Kao Kub Pla（网红泰餐，人均400泰铢）',
        '15:00 四面佛（免费，香火钱随意）',
        '16:00 暹罗百丽宫（高端购物）',
        '19:00 Jodd Fairs夜市（新火车夜市，火山排骨必吃）'
    ])
    
    add_paragraph_custom(doc, '【Day 3：文化体验日】', bold=True, size=11)
    add_bullet_list(doc, [
        '08:00 美功铁道市场（看火车穿过市场，距市区1.5小时）',
        '10:00 丹嫩沙多水上市场（坐船逛市场，船票150泰铢）',
        '14:00 返回市区，逛Terminal 21（航站楼主题商场）',
        '17:00 像素大厦观景台（门票880泰铢，看日落）',
        '20:00 唐人街夜市（燕窝、鱼翅、海鲜大排档）'
    ])
    
    # ===== 清迈3日行程 =====
    add_heading_custom(doc, '1.2 清迈3日详细行程', level=2)
    
    add_paragraph_custom(doc, '【Day 1：古城寺庙日】', bold=True, size=11)
    add_bullet_list(doc, [
        '08:00 酒店早餐',
        '09:00 帕辛寺（古城最大寺庙，免费）',
        '10:30 契迪龙寺（门票40泰铢，看大佛塔）',
        '12:00 午餐：Khao Soi Khun Yai（咖喱面，人均60泰铢）',
        '14:00 塔佩门（喂鸽子，拍照）',
        '15:00 清曼寺（最古老寺庙，免费）',
        '18:00 周日夜市（仅周日开放，世界十大夜市之一）'
    ])
    
    add_paragraph_custom(doc, '【Day 2：素贴山+宁曼路】', bold=True, size=11)
    add_bullet_list(doc, [
        '08:00 素贴山双龙寺（门票30泰铢，看全城景）',
        '11:00 蒲屏皇宫（门票50泰铢，皇家花园）',
        '14:00 宁曼路（文艺街区，咖啡店打卡）',
        '15:00 Maya商场（宁曼路地标）',
        '17:00 清迈大学（静心湖看日落）',
        '19:00 宁曼路夜市（小资夜市）'
    ])
    
    add_paragraph_custom(doc, '【Day 3：大象营+夜间动物园】', bold=True, size=11)
    add_bullet_list(doc, [
        '08:00 Elephant Nature Park（大象保护营，半天1500泰铢，不骑象）',
        '14:00 返回市区午餐',
        '16:00 清迈夜间动物园（门票800泰铢，16:00-22:00）',
        '20:00 长康路夜市（平日开放，买手信）'
    ])
    
    # ===== 普吉岛3日行程 =====
    add_heading_custom(doc, '1.3 普吉岛3日详细行程', level=2)
    
    add_paragraph_custom(doc, '【Day 1：芭东海滩日】', bold=True, size=11)
    add_bullet_list(doc, [
        '09:00 芭东海滩（免费，水上项目：滑翔伞1500泰铢/15分钟）',
        '12:00 午餐：No.6 Restaurant（网红餐厅，人均300泰铢）',
        '14:00 江西冷商场（芭东最大商场，买伴手礼）',
        '17:00 西蒙人妖秀（门票800泰铢，1小时）',
        '19:00 芭东夜市（海鲜大排档，龙虾1200泰铢/只）',
        '21:00 芭东酒吧街（感受夜生活）'
    ])
    
    add_paragraph_custom(doc, '【Day 2：跳岛游】', bold=True, size=11)
    add_bullet_list(doc, [
        '07:30 码头集合（皮皮岛一日游，1500-2500泰铢/人）',
        '09:30 玛雅湾（电影《海滩》取景地）',
        '11:00 猴子海滩',
        '12:00 大皮皮岛午餐（自助）',
        '14:00 浮潜（看珊瑚和鱼）',
        '15:30 蛋岛（喂鱼）',
        '17:00 返回普吉岛'
    ])
    
    add_paragraph_custom(doc, '【Day 3：环岛+神仙半岛】', bold=True, size=11)
    add_bullet_list(doc, [
        '09:00 查龙寺（普吉最大寺庙，免费）',
        '11:00 山顶大佛（门票免费， donations随意）',
        '13:00 午餐：悬崖餐厅（看海景，人均500泰铢）',
        '15:00 卡伦海滩（比芭东安静，水更清）',
        '17:00 神仙半岛（看日落最佳地点，免费）',
        '19:00 拉威海鲜市场（现买现做，比芭东便宜30%）'
    ])
    
    doc.add_page_break()
    
    # ===== 餐厅推荐 =====
    add_heading_custom(doc, '1.4 曼谷餐厅推荐（附地址+人均）', level=2)
    
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '餐厅名'
    hdr[1].text = '特色菜'
    hdr[2].text = '地址/区域'
    hdr[3].text = '人均'
    
    restaurants_bkk = [
        ['Thip Samai', '鬼门炒粉', '金山寺附近', '150泰铢'],
        ['Kub Kao Kub Pla', '泰式家常菜', 'Central World', '400泰铢'],
        ['Jay Fai', '米其林街头小吃', '考山路附近', '800泰铢'],
        ['Nara Thai', '精致泰餐', 'Central World', '600泰铢'],
        ['Somboon Seafood', '咖喱蟹', 'Surawong店', '800泰铢'],
        ['Mango Tango', '芒果糯米饭', 'Siam Square', '200泰铢'],
        ['After You', '蜜糖吐司', '多处分店', '300泰铢'],
        ['Boat Noodle Alley', '船面', '胜利纪念碑', '60泰铢'],
        ['T&K Seafood', '唐人街海鲜', '唐人街', '500泰铢'],
        ['Cabbages & Condoms', '创意泰餐', 'Sukhumvit', '700泰铢']
    ]
    for i, row_data in enumerate(restaurants_bkk, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    add_paragraph_custom(doc, '【清迈餐厅推荐】', bold=True, size=11)
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '餐厅名'
    hdr[1].text = '特色菜'
    hdr[2].text = '地址/区域'
    hdr[3].text = '人均'
    
    restaurants_cnx = [
        ['Khao Soi Khun Yai', '咖喱面', '古城北门', '60泰铢'],
        ['Tong Tem Toh', '泰北菜', '宁曼路', '200泰铢'],
        ['Khao Soi Islam', '牛肉咖喱面', '古城内', '80泰铢'],
        ['Dash! Restaurant', '泰式西餐', '古城内', '400泰铢'],
        ['The Riverside', '河畔餐厅', 'Ping River', '500泰铢'],
        ['Cooking Love', '泰餐', '古城内', '250泰铢'],
        ['Lert Ros', '烤鱼', '塔佩门附近', '300泰铢']
    ]
    for i, row_data in enumerate(restaurants_cnx, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ===== 酒店推荐 =====
    add_heading_custom(doc, '1.5 酒店推荐（分区域分价位）', level=2)
    
    add_paragraph_custom(doc, '【曼谷酒店】', bold=True, size=11)
    table = doc.add_table(rows=10, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '酒店名'
    hdr[1].text = '区域'
    hdr[2].text = '价位/晚'
    hdr[3].text = '特点'
    
    hotels_bkk = [
        ['The Berkeley Hotel', '暹罗', '400-500元', '位置绝佳，步行到商场'],
        ['Centre Point Silom', '是隆', '350-450元', '公寓式，带厨房'],
        ['ibis Bangkok Siam', '暹罗', '250-350元', '经济型，交通便利'],
        ['Chatrium Hotel', '河畔', '500-700元', '河景房，无边泳池'],
        ['The Okura Prestige', '奇隆', '1200-1500元', '五星，无边泳池'],
        ['Mandarin Oriental', '河畔', '2000元以上', '传奇酒店，服务顶级'],
        ['Lub d Bangkok', '是隆', '80-120元', '青旅，背包客首选'],
        ['Siam@Siam Design Hotel', '暹罗', '500-700元', '设计酒店，天台酒吧'],
        ['Anantara Siam', '奇隆', '1500-2000元', '泰式风格，花园泳池']
    ]
    for i, row_data in enumerate(hotels_bkk, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    add_paragraph_custom(doc, '【清迈酒店】', bold=True, size=11)
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '酒店名'
    hdr[1].text = '区域'
    hdr[2].text = '价位/晚'
    hdr[3].text = '特点'
    
    hotels_cnx = [
        ['De Naga Hotel', '古城', '400-600元', '兰纳风格，位置好'],
        ['U Nimman', '宁曼路', '600-800元', '设计感，年轻人喜欢'],
        ['Buri Siri Hotel', '宁曼路', '300-400元', '性价比高，泳池'],
        ['B2 Premier Hotel', '古城外', '150-250元', '经济型，干净'],
        ['137 Pillars House', '古城外', '2000元以上', '奢华精品酒店'],
        ['Deejai Backpackers', '古城', '50-80元', '青旅，氛围好']
    ]
    for i, row_data in enumerate(hotels_cnx, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    add_paragraph_custom(doc, '【普吉岛酒店】', bold=True, size=11)
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '酒店名'
    hdr[1].text = '区域'
    hdr[2].text = '价位/晚'
    hdr[3].text = '特点'
    
    hotels_hkt = [
        ['The Nai Harn', '奈汉海滩', '800-1200元', '五星，私人海滩'],
        ['Holiday Inn Resort', '芭东', '500-700元', '亲子酒店，泳池多'],
        ['The Kee Resort', '芭东', '400-600元', '位置好，近海滩'],
        ['SALA Phuket', '迈考海滩', '1000-1500元', '设计酒店，无边泳池'],
        ['Lub d Phuket', '芭东', '100-150元', '青旅，天台泳池'],
        ['Amari Phuket', '芭东', '700-1000元', '海景房，私人沙滩']
    ]
    for i, row_data in enumerate(hotels_hkt, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ===== 交通价格表 =====
    add_heading_custom(doc, '1.6 泰国交通价格表（2025年最新）', level=2)
    
    add_paragraph_custom(doc, '【曼谷交通】', bold=True, size=11)
    table = doc.add_table(rows=10, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '交通方式'
    hdr[1].text = '路线/说明'
    hdr[2].text = '价格'
    
    transport_bkk = [
        ['BTS（天铁）', '按距离计费', '16-59泰铢'],
        ['MRT（地铁）', '按距离计费', '16-42泰铢'],
        ['机场快线', '素万那普-市中心', '15-45泰铢'],
        ['出租车', '打表起步价', '35泰铢+每公里5-8泰铢'],
        ['Grab', '网约车', '比出租车贵20%'],
        ['嘟嘟车', '需议价', '100-300泰铢/次'],
        ['摩托车出租', '按天', '200-300泰铢/天'],
        ['公交船', '湄南河', '10-30泰铢'],
        ['机场大巴', '廊曼机场-市中心', '30-50泰铢']
    ]
    for i, row_data in enumerate(transport_bkk, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    add_paragraph_custom(doc, '【城际交通】', bold=True, size=11)
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '路线'
    hdr[1].text = '飞机'
    hdr[2].text = '火车'
    hdr[3].text = '大巴'
    
    intercity = [
        ['曼谷-清迈', '1500-3000泰铢/1小时', '600-1500泰铢/12小时', '500-800泰铢/10小时'],
        ['曼谷-普吉', '1200-4000泰铢/1.5小时', '-', '600-1000泰铢/12小时'],
        ['曼谷-芭提雅', '-', '-', '130-160泰铢/2小时'],
        ['曼谷-苏梅', '2500-5000泰铢/1小时', '-', '车船联运800泰铢/12小时'],
        ['清迈-普吉', '2000-4000泰铢/2小时', '-', '-'],
        ['普吉-苏梅', '2500-4000泰铢/1小时', '-', '船票600泰铢/2小时']
    ]
    for i, row_data in enumerate(intercity, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ===== 避坑清单 =====
    add_heading_custom(doc, '1.7 泰国避坑清单（必看！）', level=2)
    
    add_paragraph_custom(doc, '【交通避坑】', bold=True, size=11, color=RGBColor(204, 0, 0))
    add_bullet_list(doc, [
        '❌ 机场出租车：不要坐"Taxi Service"柜台的车，贵一倍。去官方出租车排队处， insist on meter（打表）',
        '❌ 嘟嘟车：一定要提前谈好价格，用Google Maps查距离，3公里内100-150泰铢合理',
        '❌ 摩托车出租：没驾照被警察抓到罚500泰铢，且没有保险',
        '✅ 正确做法：下载Grab，价格透明，可绑定支付宝'
    ])
    
    add_paragraph_custom(doc, '【景点避坑】', bold=True, size=11, color=RGBColor(204, 0, 0))
    add_bullet_list(doc, [
        '❌ 大皇宫门口：有人说"今天不开门"，骗你去其他地方，不要信',
        '❌ 玉佛寺：有人给你喂鸽子的玉米，接了就收费100泰铢',
        '❌ 水上市场：船夫带你去"宝石店"，价格虚高10倍',
        '❌ 人妖秀：街头拉客的秀，质量差还贵，去正规的西蒙秀或蒂芬妮秀',
        '✅ 正确做法：景点只去官方入口，不跟陌生人走，不随便接东西'
    ])
    
    add_paragraph_custom(doc, '【购物避坑】', bold=True, size=11, color=RGBColor(204, 0, 0))
    add_bullet_list(doc, [
        '❌ 珠宝/宝石：导游带去的店，价格虚高，不要买',
        '❌ 乳胶枕：超过1000泰铢就是坑，Big C超市才500泰铢',
        '❌ 蛇药/燕窝：功效夸大，价格虚高',
        '❌ 换汇：机场汇率最差，Super Rich或Value Plus最划算',
        '✅ 正确做法：购物去商场或超市，不在导游带去的店买'
    ])
    
    add_paragraph_custom(doc, '【餐饮避坑】', bold=True, size=11, color=RGBColor(204, 0, 0))
    add_bullet_list(doc, [
        '❌ 海鲜：问清楚是"每100克"还是"每公斤"价格，芭东有些店用100克报价',
        '❌ 冰沙/果汁：加冰的饮料可能用自来水冰块，肠胃不好的慎喝',
        '❌ 街头烧烤：卫生条件参差，肠胃敏感的慎吃',
        '✅ 正确做法：海鲜自己选活的，称重前确认价格单位，去人多的店'
    ])
    
    doc.add_page_break()
    
    # ===== 购物清单 =====
    add_heading_custom(doc, '1.8 泰国购物清单（买什么+去哪买+价格）', level=2)
    
    table = doc.add_table(rows=13, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '商品'
    hdr[1].text = '推荐品牌/说明'
    hdr[2].text = '去哪买'
    hdr[3].text = '参考价格'
    
    shopping = [
        ['青草膏', '卧佛牌、金杯牌', '7-11、Boots', '50-150泰铢'],
        ['鼻通', '八仙筒', '7-11', '20泰铢'],
        ['榴莲干', '金枕头', 'Big C、超市', '150-300泰铢'],
        ['芒果干', '无添加糖', 'Big C、市场', '100-200泰铢'],
        ['乳胶枕', 'Ventry、Patex', 'Big C、商场', '500-1500泰铢'],
        ['泰丝', 'Jim Thompson', '专卖店、机场', '1000-5000泰铢'],
        ['曼谷包', 'NaRaYa', '专卖店', '100-500泰铢'],
        ['欧莱雅', '全系列比国内便宜', 'Boots、屈臣氏', '国内5-7折'],
        ['牛奶洗面奶', 'Beauty Buffet', 'Boots', '130泰铢'],
        ['蛇牌爽身粉', '经典款', '7-11', '40泰铢'],
        ['虎牌膏药', '温热/清凉', '药店', '60-150泰铢'],
        ['椰子油', '护发护肤', '超市', '80-150泰铢']
    ]
    for i, row_data in enumerate(shopping, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    add_paragraph_custom(doc, '【退税说明】', bold=True, size=11)
    add_bullet_list(doc, [
        '同一商场满2000泰铢可退税，退7%',
        '保留小票，去商场退税柜台开退税单',
        '机场退税：先盖章（Check-in前），再拿钱（过安检后）',
        '奢侈品（珠宝、手表）需随身携带检查'
    ])
    
    doc.add_page_break()
    
    # ===== 实用泰语 =====
    add_heading_custom(doc, '1.9 实用泰语（附发音）', level=2)
    
    add_paragraph_custom(doc, '【基础问候】', bold=True, size=11)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '中文'
    hdr[1].text = '泰语'
    hdr[2].text = '发音'
    
    thai_basic = [
        ['你好', 'สวัสดี', '萨瓦迪卡'],
        ['谢谢', 'ขอบคุณ', '阔坤卡'],
        ['对不起', 'ขอโทษ', '阔托'],
        ['多少钱', 'เท่าไหร่', '套来'],
        ['太贵了', 'แพงมาก', '啪马克']
    ]
    for i, row_data in enumerate(thai_basic, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    add_paragraph_custom(doc, '【交通用语】', bold=True, size=11)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '中文'
    hdr[1].text = '泰语'
    hdr[2].text = '发音'
    
    thai_transport = [
        ['打表', 'เมเตอร์', '咩特'],
        ['去这里', 'ไปที่นี่', '拜替尼'],
        ['多少钱', 'เท่าไหร่', '套来'],
        ['便宜点', 'ถูกกว่านี้', '图夸尼'],
        ['停车', 'จอดตรงนี้', '作东尼'],
        ['机场', 'สนามบิน', '萨南宾']
    ]
    for i, row_data in enumerate(thai_transport, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    add_paragraph_custom(doc, '【餐饮用语】', bold=True, size=11)
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '中文'
    hdr[1].text = '泰语'
    hdr[2].text = '发音'
    
    thai_food = [
        ['好吃', 'อร่อย', '阿来'],
        ['辣', 'เผ็ด', '佩'],
        ['不辣', 'ไม่เผ็ด', '迈佩'],
        ['买单', 'เช็คบิล', '切宾'],
        ['水', 'น้ำ', '南'],
        ['冰', 'น้ำแข็ง', '南康'],
        ['米饭', 'ข้าว', '靠']
    ]
    for i, row_data in enumerate(thai_food, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ==================== 日本 ====================
    add_heading_custom(doc, '四、日本完全指南', level=1)
    
    add_heading_custom(doc, '4.1 东京5日详细行程', level=2)
    
    add_paragraph_custom(doc, '【Day 1：浅草-晴空塔-秋叶原】', bold=True, size=11)
    add_bullet_list(doc, [
        '08:00 酒店出发',
        '09:00 浅草寺（免费，求签100日元，雷门拍照）',
        '11:00 仲见世通（买人形烧，150日元/个）',
        '12:00 午餐：大黑家天妇罗（人均1500日元）',
        '14:00 晴空塔（350米观景台2100日元，450米3100日元）',
        '17:00 秋叶原（电器街，动漫周边，女仆咖啡厅）',
        '19:00 晚餐：矶丸水产（海鲜烧烤，人均3000日元）'
    ])
    
    add_paragraph_custom(doc, '【Day 2：涩谷-原宿-新宿】', bold=True, size=11)
    add_bullet_list(doc, [
        '09:00 涩谷十字路口（世界上最忙的十字路口）',
        '10:00 涩谷Sky（观景台2200日元，需预约）',
        '12:00 午餐：一兰拉面（人均1200日元）',
        '14:00 原宿竹下通（年轻人文化，可丽饼300日元）',
        '16:00 明治神宫（免费，森林中的神社）',
        '18:00 新宿（哥斯拉酒店拍照）',
        '19:00 晚餐：思い出横丁（烤鸡串一条街，人均2500日元）',
        '21:00 歌舞伎町（感受夜生活，注意安全）'
    ])
    
    add_paragraph_custom(doc, '【Day 3：富士山一日游】', bold=True, size=11)
    add_bullet_list(doc, [
        '07:30 新宿巴士总站出发（往返4000日元，2小时）',
        '10:00 河口湖（看富士山倒影，缆车800日元）',
        '12:00 午餐：ほうとう不動（乌冬面，人均1200日元）',
        '14:00 忍野八海（免费，富士山雪水形成的清泉）',
        '16:00 御殿场奥特莱斯（购物，品牌折扣）',
        '19:00 返回东京'
    ])
    
    add_paragraph_custom(doc, '【Day 4：镰仓一日游】', bold=True, size=11)
    add_bullet_list(doc, [
        '08:00 新宿出发，坐小田急线（江之电一日券800日元）',
        '10:00 镰仓大佛（门票300日元）',
        '11:30 长谷寺（门票400日元，看紫阳花/海景）',
        '13:00 午餐：鎌倉釜飯かまかま（釜饭，人均1500日元）',
        '15:00 镰仓高校前站（灌篮高手打卡地）',
        '16:00 江之岛（海边散步，灯塔500日元）',
        '18:00 返回东京'
    ])
    
    add_paragraph_custom(doc, '【Day 5：筑地-银座-表参道】', bold=True, size=11)
    add_bullet_list(doc, [
        '07:00 筑地场外市场（寿司大，排队1小时起，人均4000日元）',
        '10:00 银座（奢侈品一条街， window shopping）',
        '12:00 午餐：银座九丁目（天妇罗，人均2000日元）',
        '14:00 表参道（潮牌一条街，建筑好看）',
        '16:00 根津美术馆（门票1300日元，庭院很美）',
        '18:00 六本木之丘（看东京塔夜景，门票1800日元）',
        '20:00 晚餐：叙々苑（烤肉，人均5000日元）'
    ])
    
    doc.add_page_break()
    
    add_heading_custom(doc, '4.2 日本餐厅推荐（附地址+人均）', level=2)
    
    table = doc.add_table(rows=12, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '餐厅名'
    hdr[1].text = '特色'
    hdr[2].text = '区域'
    hdr[3].text = '人均'
    
    restaurants_jp = [
        ['寿司大', ' omakase寿司', '筑地', '4000日元'],
        ['一兰拉面', '豚骨拉面', '多处分店', '1200日元'],
        ['鸟贵族', '烤鸡串', '多处分店', '3000日元'],
        ['蟹道乐', '螃蟹料理', '新宿/银座', '8000日元'],
        ['叙々苑', '和牛烤肉', '新宿', '5000日元'],
        ['大黑家', '天妇罗', '浅草', '1500日元'],
        ['矶丸水产', '海鲜烧烤', '多处分店', '3000日元'],
        ['松屋/吉野家', '牛肉饭', '多处分店', '500日元'],
        ['HARBS', '水果千层蛋糕', '表参道', '1500日元'],
        ['中村藤吉', '抹茶甜品', '浅草', '1000日元'],
        ['Luke\'s Lobster', '龙虾卷', '原宿', '1500日元']
    ]
    for i, row_data in enumerate(restaurants_jp, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    add_heading_custom(doc, '4.3 日本交通价格表', level=2)
    
    table = doc.add_table(rows=10, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '交通方式'
    hdr[1].text = '说明'
    hdr[2].text = '价格'
    
    transport_jp = [
        ['JR Pass（全国）', '7日券', '约1500元'],
        ['JR Pass（关东）', '3日券', '约600元'],
        ['JR Pass（关西）', '3日券', '约550元'],
        ['东京地铁24小时券', '都营+Metro', '800日元'],
        ['东京地铁48小时券', '都营+Metro', '1200日元'],
        ['西瓜卡(Suica)', '储值卡', '最低充值1000日元'],
        ['新干线东京-京都', '约2.5小时', '约700元'],
        ['新干线东京-大阪', '约2.5小时', '约750元'],
        ['成田机场-市区', 'N\'EX/京成Skyliner', '约150-250元']
    ]
    for i, row_data in enumerate(transport_jp, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    add_heading_custom(doc, '4.4 日本避坑清单', level=2)
    add_bullet_list(doc, [
        '❌ 不要在机场换太多现金，汇率差，带银联卡直接取现最划算',
        '❌ 不要在日本打车，贵到离谱（起步费约50元，跳表很快）',
        '❌ 不要在便利店买水果，一个桃子可能100元，去超市或市场',
        '❌ 不要边走边吃，在便利店买的东西要在店里吃完再走',
        '❌ 不要大声说话，电车内不要打电话',
        '❌ 不要给小费，日本没有小费文化',
        '✅ 正确做法：提前买JR Pass，用Google Maps查路线，便利店买早餐'
    ])
    
    doc.add_page_break()
    
    # ==================== 万能行李清单 ====================
    add_heading_custom(doc, '六、万能行李清单（所有目的地通用）', level=1)
    
    add_heading_custom(doc, '6.1 证件类', level=2)
    add_bullet_list(doc, [
        '护照（有效期6个月以上）',
        '签证（如需）',
        '机票行程单（电子版即可）',
        '酒店预订单（电子版即可）',
        '身份证（国内转机用）',
        '驾照（如需租车，带翻译件）',
        '保险单（电子版）'
    ])
    
    add_heading_custom(doc, '6.2 电器类', level=2)
    add_bullet_list(doc, [
        '手机+充电器',
        '充电宝（不超过100Wh，随身携带）',
        '转换插头（日本/韩国两孔通用，泰国可直接用，新加坡需三孔转换）',
        '插线板（酒店插座不够用）',
        '耳机',
        '相机（可选）'
    ])
    
    add_heading_custom(doc, '6.3 衣物类', level=2)
    add_bullet_list(doc, [
        '内衣裤（按天数+2套）',
        '袜子',
        'T恤/上衣',
        '长裤/长裙（寺庙需要）',
        '短裤',
        '薄外套（空调房/飞机上冷）',
        '泳衣（海边用）',
        '舒适的 walking shoes',
        '拖鞋（酒店/沙滩）',
        '帽子+墨镜',
        '防晒衣'
    ])
    
    add_heading_custom(doc, '6.4 洗漱用品', level=2)
    add_bullet_list(doc, [
        '牙刷+牙膏（部分酒店不提供）',
        '洗发水+沐浴露（旅行装）',
        '洗面奶',
        '护肤品',
        '防晒霜（必须！）',
        '化妆品',
        '剃须刀',
        '毛巾（可带一次性）',
        '纸巾+湿巾'
    ])
    
    add_heading_custom(doc, '6.5 药品类', level=2)
    add_bullet_list(doc, [
        '肠胃药（水土不服）',
        '感冒药',
        '创可贴',
        '晕车药',
        '止痛药',
        '防蚊液（东南亚必备）',
        '个人常用药',
        '口罩'
    ])
    
    add_heading_custom(doc, '6.6 其他', level=2)
    add_bullet_list(doc, [
        '雨伞/雨衣',
        '水杯',
        '零食（飞机上吃）',
        '笔记本+笔（填入境卡）',
        '颈枕（长途飞行）',
        '眼罩+耳塞',
        '密封袋（装湿衣服）',
        '现金（提前换好）',
        '信用卡（Visa/Mastercard）'
    ])
    
    doc.add_page_break()
    
    # ==================== 紧急联系 ====================
    add_heading_custom(doc, '七、紧急联系大全', level=1)
    
    add_heading_custom(doc, '7.1 中国领事保护', level=2)
    add_paragraph_custom(doc, '全球领事保护热线：+86-10-12308 或 +86-10-59913991（24小时）')
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '国家'
    hdr[1].text = '领事馆电话'
    hdr[2].text = '地址'
    
    embassies = [
        ['泰国', '+66-2-245-7044', '曼谷拉差达路'],
        ['新加坡', '+65-6475-0165', '东陵路150号'],
        ['马来西亚', '+60-3-2164-5301', '吉隆坡安邦路'],
        ['日本', '+81-3-6450-2195', '东京都港区'],
        ['韩国', '+82-2-755-0572', '首尔市中区']
    ]
    for i, row_data in enumerate(embassies, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    add_heading_custom(doc, '7.2 当地紧急电话', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '国家'
    hdr[1].text = '报警/急救'
    hdr[2].text = '旅游警察'
    
    emergency = [
        ['泰国', '191', '1155（中文）'],
        ['新加坡', '999', '无'],
        ['马来西亚', '999', '无'],
        ['日本', '110（报警）/119（急救）', '无'],
        ['韩国', '112', '1330（中文）']
    ]
    for i, row_data in enumerate(emergency, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
    
    add_heading_custom(doc, '7.3 其他重要电话', level=2)
    add_bullet_list(doc, [
        '信用卡挂失：Visa +1-303-967-1096，Mastercard +1-636-722-7111',
        '支付宝海外：+86-571-2688-8888',
        '微信海外：+86-571-2688-8888',
        '国际机票改签：联系原购买渠道'
    ])
    
    # 结尾
    doc.add_paragraph()
    doc.add_paragraph()
    ending = doc.add_paragraph()
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ending.add_run('祝你旅途愉快！\n\n本攻略由 AI 旅行规划工具 整理制作')
    run.font.name = '微软雅黑'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    return doc

if __name__ == '__main__':
    doc = create_travel_guide_v2()
    output_path = r'C:/Users/hua/Desktop/旅游攻略.docx'
    doc.save(output_path)
    print(f'文档已保存到: {output_path}')
