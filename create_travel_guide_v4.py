#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
东南亚+日韩旅游攻略生成器 v4
7个国家完整版：泰国、越南、马来西亚、新加坡、印尼、日本、韩国
优化：表格呈现 + 详细目录
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    from docx.oxml import parse_xml
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_obj = tcBorders.find(qn(f'w:{edge}'))
            if edge_obj is None:
                edge_obj = parse_xml(f'<w:{edge} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                tcBorders.append(edge_obj)
            edge_obj.set(qn('w:val'), kwargs[edge])

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 102, 204)
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 153, 76)
        run.font.bold = True
    else:
        run.font.size = Pt(12)
        run.font.bold = True
    return heading

def add_table_from_data(doc, headers, data, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # 添加数据行
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    return table

def create_travel_guide():
    doc = Document()
    
    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    
    # 封面
    title = doc.add_heading('东南亚 + 日韩旅游完全攻略', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 102, 204)
        run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('泰国 · 越南 · 马来西亚 · 新加坡 · 印尼 · 日本 · 韩国\n实用工具书 · 可直接查询使用')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    
    # ========== 详细目录 ==========
    add_heading_custom(doc, '目录', level=1)
    
    toc_sections = [
        ('一、泰国攻略', [
            '1.1 核心城市与景点（曼谷/清迈/普吉/芭提雅）',
            '1.2 餐厅推荐（分区域/分价位）',
            '1.3 酒店推荐（分区域/分价位）',
            '1.4 交通价格参考',
            '1.5 避坑清单',
            '1.6 购物清单',
            '1.7 实用泰语'
        ]),
        ('二、越南攻略', [
            '2.1 核心城市与景点（河内/岘港/芽庄/胡志明）',
            '2.2 餐厅推荐',
            '2.3 酒店推荐',
            '2.4 交通价格',
            '2.5 避坑清单',
            '2.6 购物清单',
            '2.7 实用越南语'
        ]),
        ('三、马来西亚攻略', [
            '3.1 核心城市与景点（吉隆坡/槟城/兰卡威）',
            '3.2 餐厅推荐',
            '3.3 酒店推荐',
            '3.4 交通价格',
            '3.5 避坑清单',
            '3.6 购物清单'
        ]),
        ('四、新加坡攻略', [
            '4.1 核心景点',
            '4.2 餐厅推荐',
            '4.3 酒店推荐',
            '4.4 交通价格',
            '4.5 避坑清单',
            '4.6 购物清单'
        ]),
        ('五、印尼攻略', [
            '5.1 核心目的地（巴厘岛/雅加达）',
            '5.2 餐厅推荐',
            '5.3 酒店推荐',
            '5.4 交通价格',
            '5.5 避坑清单',
            '5.6 购物清单'
        ]),
        ('六、日本攻略', [
            '6.1 核心城市与景点（东京/京都/大阪）',
            '6.2 餐厅推荐',
            '6.3 酒店推荐',
            '6.4 交通价格',
            '6.5 避坑清单',
            '6.6 购物清单',
            '6.7 实用日语'
        ]),
        ('七、韩国攻略', [
            '7.1 核心城市与景点（首尔/釜山/济州岛）',
            '7.2 餐厅推荐',
            '7.3 酒店推荐',
            '7.4 交通价格',
            '7.5 避坑清单',
            '7.6 购物清单',
            '7.7 实用韩语'
        ]),
        ('八、通用工具', [
            '8.1 万能行李清单',
            '8.2 紧急联系电话',
            '8.3 汇率速查表'
        ])
    ]
    
    for section, subsections in toc_sections:
        p = doc.add_paragraph()
        p.add_run(section).bold = True
        for sub in subsections:
            p = doc.add_paragraph(sub, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # ==================== 泰国 ====================
    add_heading_custom(doc, '一、泰国攻略', level=1)
    
    add_heading_custom(doc, '1.1 核心城市与景点', level=2)
    
    # 曼谷景点表格
    add_heading_custom(doc, '【曼谷】- 寺庙与现代交融', level=3)
    doc.add_paragraph()
    
    bkk_headers = ['景点名称', '门票', '开放时间', '特色/备注']
    bkk_data = [
        ['大皇宫+玉佛寺', '500泰铢', '8:30-15:30', '着装要求：长裤长裙，不能露肩'],
        ['卧佛寺', '200泰铢', '8:00-18:30', '曼谷最大寺庙，泰式按摩发源地'],
        ['郑王庙（黎明寺）', '100泰铢', '8:00-18:00', '网红拍照点，可俯瞰湄南河'],
        ['四面佛', '免费', '6:00-22:00', '位于Central World旁，香火最旺'],
        ['水门寺大佛', '免费', '8:00-18:00', '2022新建，69米高金色大佛'],
        ['ICONSIAM商场', '免费', '10:00-22:00', '室内水上市场，高端购物'],
        ['Jodd Fairs夜市', '免费', '17:00-00:00', '网红火山排骨、水果西施'],
        ['Chatuchak周末市场', '免费', '周六日9:00-18:00', '亚洲最大周末市场']
    ]
    add_table_from_data(doc, bkk_headers, bkk_data)
    doc.add_paragraph()
    
    # 清迈景点表格
    add_heading_custom(doc, '【清迈】- 古城与慢生活', level=3)
    doc.add_paragraph()
    
    cnx_headers = ['景点名称', '门票', '开放时间', '特色/备注']
    cnx_data = [
        ['塔佩门', '免费', '全天', '古城东门，喂鸽子拍照，周日夜市起点'],
        ['双龙寺（素贴山）', '30泰铢', '6:00-20:00', '俯瞰清迈全城，金色佛塔'],
        ['宁曼路', '免费', '全天', '文青聚集地，咖啡馆、买手店'],
        ['周日夜市', '免费', '周日16:00-22:00', '古城主干道，手工艺品、小吃'],
        ['大象保护营', '1500-2500泰铢', '半天/全天', '给大象洗澡喂食，无表演'],
        ['夜间动物园', '800泰铢', '17:00-22:00', '世界最大夜间动物园'],
        ['丛林飞跃', '1500-3000泰铢', '全天', 'Skyline/Flight of Gibbon，最长900米'],
        ['拜县', '免费', '全天', '童话小镇，需2天，黄色小屋、大树秋千']
    ]
    add_table_from_data(doc, cnx_headers, cnx_data)
    doc.add_paragraph()
    
    # 普吉景点表格
    add_heading_custom(doc, '【普吉岛】- 海岛度假首选', level=3)
    doc.add_paragraph()
    
    hkt_headers = ['景点名称', '门票/价格', '特色/备注']
    hkt_data = [
        ['芭东海滩', '免费', '最热闹海滩，酒吧街、江西冷商场'],
        ['卡伦/卡塔海滩', '免费', '水质更好，适合冲浪、看日落'],
        ['皮皮岛一日游', '1200-2500泰铢', '玛雅湾、猴子沙滩、浮潜'],
        ['皇帝岛/珊瑚岛', '1500-3000泰铢', '水质清澈，深潜圣地'],
        ['西蒙人妖秀', '800泰铢', '18:00/19:30/21:00，最正规人妖秀'],
        ['普吉老镇', '免费', '彩色葡式建筑，米其林餐厅聚集'],
        ['神仙半岛', '免费', '最佳日落观景点'],
        ['查龙寺', '免费', '普吉最大寺庙，可放鞭炮祈福']
    ]
    add_table_from_data(doc, hkt_headers, hkt_data)
    doc.add_paragraph()
    
    # 芭提雅景点表格
    add_heading_custom(doc, '【芭提雅】- 夜生活与海岛', level=3)
    doc.add_paragraph()
    
    pattaya_headers = ['景点名称', '门票/价格', '特色/备注']
    pattaya_data = [
        ['蒂芬妮人妖秀', '1200泰铢', '18:00/19:30/21:00，泰国最著名人妖秀'],
        ['格兰岛一日游', '500-1500泰铢', '水上项目齐全，拖拽伞、香蕉船'],
        ['真理寺', '500泰铢', '8:00-18:00，全木雕建筑，壮观'],
        ['四方水上市场', '200泰铢', '9:00-20:00，《杜拉拉升职记》取景地'],
        ['风月步行街', '免费', '18:00-02:00，酒吧、夜店聚集']
    ]
    add_table_from_data(doc, pattaya_headers, pattaya_data)
    doc.add_paragraph()
    
    # 餐厅推荐表格
    add_heading_custom(doc, '1.2 餐厅推荐（分区域/分价位）', level=2)
    doc.add_paragraph()
    
    restaurant_headers = ['城市', '餐厅名称', '价位', '特色/推荐菜', '备注']
    restaurant_data = [
        ['曼谷', 'Jay Fai（痣姐热炒）', '150-250泰铢', '蟹肉煎蛋卷、冬阴功', '米其林街头美食，需排队1-2小时'],
        ['曼谷', 'Thip Samai（鬼门炒粉）', '100-150泰铢', '泰式炒河粉', '16:00-02:00'],
        ['曼谷', 'Krua Apsorn', '200-300泰铢', '蟹肉煎蛋、冬阴功', '皇室认证泰餐'],
        ['曼谷', '水门海南鸡饭', '60-80泰铢', '海南鸡饭', '红/绿制服两家'],
        ['曼谷', 'Nara Thai Cuisine', '600-800泰铢', '环境好泰餐', 'Central World 7楼'],
        ['曼谷', 'Somboon Seafood（建兴酒家）', '800-1200泰铢', '咖喱蟹必点', '多分店'],
        ['曼谷', 'Gaggan Anand', '5000+泰铢', '创意印度菜', '亚洲最佳餐厅，需预约'],
        ['清迈', 'Khao Soi Khun Yai', '50-80泰铢', '泰北咖喱面', '仅午餐'],
        ['清迈', 'Tong Tem Toh', '200-300泰铢', '泰北菜、烤猪肉', ''],
        ['清迈', 'The House by Ginger', '400-600泰铢', '创意泰餐', '环境好'],
        ['普吉', 'No.6 Restaurant', '300-500泰铢', '泰餐', '芭东网红，需排队'],
        ['普吉', 'Kan Eang@Pier', '800-1500泰铢', '海鲜', '查龙码头'],
        ['普吉', 'Raya Restaurant', '400-600泰铢', '传统泰餐', '普吉镇']
    ]
    add_table_from_data(doc, restaurant_headers, restaurant_data)
    doc.add_paragraph()
    
    # 酒店推荐表格
    add_heading_custom(doc, '1.3 酒店推荐（分区域/分价位）', level=2)
    doc.add_paragraph()
    
    hotel_headers = ['城市/区域', '酒店名称', '价位', '特色/备注']
    hotel_data = [
        ['曼谷暹罗区', 'Lub d Bangkok Siam', '150-300元/晚', '青旅风格，近BTS'],
        ['曼谷暹罗区', 'Holiday Inn Express Siam', '400-600元/晚', '位置极佳'],
        ['曼谷暹罗区', 'Siam Kempinski Hotel', '1500-2500元/晚', '直连商场，有泳池'],
        ['曼谷素坤逸', 'The Salil Hotel Sukhumvit 57', '300-500元/晚', '网红酒店，拍照好看'],
        ['曼谷素坤逸', 'Marriott Marquis Queen\'s Park', '1000-1800元/晚', '五星级，设施完善'],
        ['曼谷湄南河', 'Mandarin Oriental Bangkok', '2500-5000元/晚', '百年传奇酒店'],
        ['曼谷湄南河', 'The Peninsula Bangkok', '3000-6000元/晚', '顶级服务'],
        ['清迈', 'De Lanna Hotel', '150-250元/晚', '古城内，泰式风格'],
        ['清迈', 'Buri Siri Boutique Hotel', '300-500元/晚', '宁曼路附近，设计感强'],
        ['清迈', '137 Pillars House', '1500-3000元/晚', '殖民风格，服务顶级'],
        ['普吉芭东', 'Lub d Phuket Patong', '200-350元/晚', '青旅+酒店'],
        ['普吉芭东', 'Amari Phuket', '1200-2000元/晚', '私人海滩'],
        ['普吉卡伦', 'Sawasdee Village', '800-1500元/晚', '泰式风格，泳池美'],
        ['普吉卡伦', 'The Shore at Katathani', '2000-4000元/晚', '仅限成人，私密性高']
    ]
    add_table_from_data(doc, hotel_headers, hotel_data)
    doc.add_paragraph()
    
    # 交通价格表格
    add_heading_custom(doc, '1.4 交通价格参考', level=2)
    doc.add_paragraph()
    
    transport_headers = ['交通方式', '价格', '备注']
    transport_data = [
        ['曼谷BTS（轻轨）', '15-59泰铢/程', '按距离计费'],
        ['曼谷MRT（地铁）', '16-42泰铢/程', ''],
        ['曼谷出租车', '35泰铢起步', '每公里5-8泰铢，一定要打表'],
        ['Grab打车', '比出租车贵20-30%', '明码标价'],
        ['嘟嘟车（Tuk-tuk）', '100-300泰铢/程', '需砍价'],
        ['摩的', '10-50泰铢/程', '堵车时最快'],
        ['曼谷-清迈机票', '800-2000泰铢', '提前订便宜'],
        ['曼谷-清迈大巴', '600-1000泰铢', 'VIP大巴约10小时'],
        ['曼谷-普吉机票', '1000-3000泰铢', ''],
        ['普吉机场-芭东海滩', '150-200泰铢（机场巴士）', '出租车800-1000泰铢'],
        ['清迈双条车', '20-50泰铢/人', '红色，包车200-500泰铢'],
        ['普吉岛租摩托车', '200-400泰铢/天', '需国际驾照或翻译件']
    ]
    add_table_from_data(doc, transport_headers, transport_data)
    doc.add_paragraph()
    
    # 避坑清单表格
    add_heading_custom(doc, '1.5 避坑清单', level=2)
    doc.add_paragraph()
    
    scam_headers = ['坑点', '具体表现', '对策']
    scam_data = [
        ['大皇宫门口骗局', '有人说"大皇宫今天不开门"', '无视，直接走到正门'],
        ['出租车不打表', '司机说"by meter broken"', '坚持打表，不打就换一辆'],
        ['嘟嘟车乱开价', '开口就要300-500泰铢', '提前查好距离，砍价到100-200'],
        ['水上市场一日游陷阱', '报价极低，但带去购物点', '选择正规旅行社或自行前往'],
        ['人妖秀拍照收费', '主动拉你拍照，拍完要小费', '不想拍就拒绝'],
        ['按摩店价格不透明', '街边店可能临时加价', '选择Let\'s Relax等连锁品牌'],
        ['租摩托车押金纠纷', '还车时以各种理由扣押金', '租车时拍照录像，确认划痕'],
        ['海鲜餐厅宰客', '芭东部分餐厅价格虚高', '看菜单明码标价，或去班赞市场买后加工']
    ]
    add_table_from_data(doc, scam_headers, scam_data)
    doc.add_paragraph()
    
    # 购物清单
    add_heading_custom(doc, '1.6 购物清单', level=2)
    doc.add_paragraph()
    
    shopping_headers = ['类别', '推荐商品', '推荐购买地点']
    shopping_data = [
        ['零食', '小老板海苔、大哥花生、Bento鱿鱼片、MAMA泡面', '7-11、Big C超市'],
        ['药品', '青草膏、蜈蚣丸、虎牌膏药、鼻通、soffell驱蚊水', 'Boots、药店'],
        ['化妆品', 'Mistine彩妆、Beauty Buffet牛奶系列、蜗牛霜', 'Boots、EVEANDBOY'],
        ['服饰', '曼谷包（NaRaYa）、泰丝、泰式服饰', 'Central World、Chatuchak'],
        ['工艺品', '佛牌、木雕、银器、香薰', 'Chatuchak、夜市'],
        ['退税', '满2000泰铢可退', 'VAT Refund标志的店，机场办理']
    ]
    add_table_from_data(doc, shopping_headers, shopping_data)
    doc.add_paragraph()
    
    # 实用泰语表格
    add_heading_custom(doc, '1.7 实用泰语', level=2)
    doc.add_paragraph()
    
    thai_headers = ['中文', '泰语', '发音']
    thai_data = [
        ['你好', 'Sa-wat-dee', '萨瓦迪卡'],
        ['谢谢', 'Khop-khun', '考坤'],
        ['对不起', 'Khor-thot', '考托'],
        ['多少钱', 'Tao-rai', '套来'],
        ['太贵了', 'Phaeng-mai', '潘卖'],
        ['便宜点', 'Lot-rai-hoi-noi', '落来诺依'],
        ['好吃', 'A-roy', '阿洛依'],
        ['厕所在哪', 'Hong-nam-yu-tee-nai', '洪南由提乃'],
        ['救命', 'Chuay-duay', '求端'],
        ['1', 'nung', 'nung'],
        ['2', 'song', 'song'],
        ['3', 'sam', 'sam']
    ]
    add_table_from_data(doc, thai_headers, thai_data)
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ==================== 越南 ====================
    add_heading_custom(doc, '二、越南攻略', level=1)
    
    add_heading_custom(doc, '2.1 核心城市与景点', level=2)
    doc.add_paragraph()
    
    vn_headers = ['城市', '景点名称', '门票/价格', '特色/备注']
    vn_data = [
        ['河内', '还剑湖', '免费', '市中心地标，早晨有当地人晨练'],
        ['河内', '老城区（36行街）', '免费', '每条街卖一种商品，摩托车穿行'],
        ['河内', '圣约瑟夫大教堂', '免费', '河内版巴黎圣母院，哥特式'],
        ['河内', '水上木偶戏', '100,000-200,000盾', 'Thang Long Water Puppet Theatre'],
        ['河内', '火车街', '免费', '火车从居民楼间穿行，网红打卡'],
        ['岘港', '美溪海滩', '免费', '世界最美六大海滩之一'],
        ['岘港', '巴拿山', '750,000盾', '法式山城，佛手桥、游乐园'],
        ['岘港', '会安古城', '免费入城', '灯笼夜景绝美，tailor定制奥黛'],
        ['芽庄', '四岛游', '200,000-400,000盾', '浮潜、水上项目'],
        ['芽庄', '泥浆浴', '150,000-300,000盾', 'I-Resort或Thap Ba'],
        ['芽庄', '婆那加占婆塔', '22,000盾', '印度教遗址，小吴哥窟'],
        ['芽庄', '珍珠岛乐园', '880,000盾', '跨海缆车、水上乐园'],
        ['胡志明', '中央邮局', '免费', '法式建筑，可寄明信片'],
        ['胡志明', '红教堂', '免费', '红砖外墙，对面就是邮局'],
        ['胡志明', '统一宫', '40,000盾', '历史建筑，了解越战'],
        ['胡志明', '战争遗迹博物馆', '40,000盾', '越战历史，震撼'],
        ['胡志明', '范五老街', '免费', '背包客聚集地，酒吧街']
    ]
    add_table_from_data(doc, vn_headers, vn_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '2.2 餐厅推荐', level=2)
    doc.add_paragraph()
    
    vn_rest_headers = ['城市', '餐厅名称', '人均', '特色/推荐菜']
    vn_rest_data = [
        ['河内', 'Pho Gia Truyen（小燕糯米饭）', '30,000-50,000盾', '糯米饭配肉'],
        ['河内', 'Bun Cha Huong Lien（奥巴马米粉）', '50,000-80,000盾', '烤肉米粉，奥巴马同款'],
        ['河内', 'Pho Thin', '40,000-60,000盾', '牛肉河粉'],
        ['河内', 'Cafe Giang', '25,000盾', '鸡蛋咖啡发源地'],
        ['岘港', 'Mi Quang 1A', '40,000-60,000盾', '广南面，岘港特色'],
        ['岘港', 'Banh Xeo Ba Duong', '30,000-50,000盾', '越南煎饼'],
        ['会安', 'Morning Glory', '200,000-400,000盾', '高档越南菜'],
        ['芽庄', 'Lac Canh Restaurant', '150,000-300,000盾', '烧烤海鲜'],
        ['芽庄', 'Nem Nuong Nha Trang', '50,000-100,000盾', '芽庄春卷']
    ]
    add_table_from_data(doc, vn_rest_headers, vn_rest_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '2.3 酒店推荐', level=2)
    doc.add_paragraph()
    
    vn_hotel_headers = ['城市', '酒店名称', '价位', '特色/备注']
    vn_hotel_data = [
        ['河内', 'Hanoi La Siesta Hotel', '400-700元/晚', '老城区，服务超好'],
        ['河内', 'Hilton Hanoi Opera', '800-1200元/晚', '法式风格'],
        ['岘港', 'Novotel Danang Premier', '500-800元/晚', '江景'],
        ['岘港', 'InterContinental Danang', '2000-4000元/晚', '顶级度假'],
        ['芽庄', 'Liberty Central Nha Trang', '300-500元/晚', '海景'],
        ['芽庄', 'Vinpearl Resort & Spa', '1000-2000元/晚', '珍珠岛上'],
        ['胡志明', 'Liberty Central Saigon Centre', '400-700元/晚', '市中心'],
        ['胡志明', 'Hotel des Arts Saigon', '1000-1800元/晚', '法式奢华']
    ]
    add_table_from_data(doc, vn_hotel_headers, vn_hotel_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '2.4 交通价格', level=2)
    doc.add_paragraph()
    
    vn_trans_headers = ['交通方式', '价格', '备注']
    vn_trans_data = [
        ['河内/胡志明出租车', '12,000-15,000盾起步', '每公里11,000-15,000盾'],
        ['Grab', '比出租车便宜', '明码标价'],
        ['摩托车出租', '100,000-200,000盾/天', ''],
        ['Open Bus（长途巴士）', '400,000-600,000盾', 'Futa、新咖啡，河内-胡志明'],
        ['越南国内航班', '800,000-1,500,000盾', '胡志明-河内'],
        ['火车', '400,000-800,000盾', '河内-岘港，卧铺']
    ]
    add_table_from_data(doc, vn_trans_headers, vn_trans_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '2.5 避坑清单', level=2)
    doc.add_paragraph()
    
    vn_scam_headers = ['坑点', '对策']
    vn_scam_data = [
        ['出租车绕路/调表', '用Grab，或提前查好路线'],
        ['找零少给', '越南盾面额大，当面数清'],
        ['摩托车抢包', '包背内侧，手机别拿手上'],
        ['购物宰客', '去大商场或超市'],
        ['海关小费', '坚决不给，装听不懂'],
        ['假咖啡', '去Trung Nguyen等正规品牌店']
    ]
    add_table_from_data(doc, vn_scam_headers, vn_scam_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '2.6 购物清单', level=2)
    doc.add_paragraph()
    
    vn_shop_headers = ['类别', '推荐商品', '推荐购买地点']
    vn_shop_data = [
        ['咖啡', 'G7速溶、Trung Nguyen咖啡豆、猫屎咖啡', '超市、咖啡店'],
        ['零食', '椰子糖、榴莲饼、腰果、蔬果干', 'Big C、Ben Thanh市场'],
        ['工艺品', '奥黛、漆器、竹编、丝绸', '会安夜市'],
        ['药品', '白虎膏、蛇油', '药店']
    ]
    add_table_from_data(doc, vn_shop_headers, vn_shop_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '2.7 实用越南语', level=2)
    doc.add_paragraph()
    
    vn_lang_headers = ['中文', '越南语', '发音']
    vn_lang_data = [
        ['你好', 'Xin chao', '新交'],
        ['谢谢', 'Cam on', '感恩'],
        ['多少钱', 'Bao nhieu tien', '包尼乌田'],
        ['太贵了', 'Dat qua', '达夸'],
        ['便宜点', 'Re hon di', '热昏迪'],
        ['好吃', 'Ngon', '嗯']
    ]
    add_table_from_data(doc, vn_lang_headers, vn_lang_data)
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ==================== 马来西亚 ====================
    add_heading_custom(doc, '三、马来西亚攻略', level=1)
    
    add_heading_custom(doc, '3.1 核心城市与景点', level=2)
    doc.add_paragraph()
    
    my_attr_headers = ['城市', '景点名称', '门票/价格', '特色/备注']
    my_attr_data = [
        ['吉隆坡', '双子塔（Petronas Towers）', '免费外观，登顶RM80', '需预约'],
        ['吉隆坡', '独立广场', '免费', '殖民建筑，拍照打卡'],
        ['吉隆坡', '茨厂街（唐人街）', '免费', '美食、购物'],
        ['吉隆坡', '黑风洞（Batu Caves）', '免费', '272级台阶，彩虹阶梯'],
        ['吉隆坡', '阿罗街（Jalan Alor）', '免费', '夜市美食街'],
        ['吉隆坡', '国家清真寺', '免费', '9:00-18:00，需穿长袍'],
        ['槟城', '乔治市壁画街', '免费', '网红壁画打卡'],
        ['槟城', '姓氏桥', '免费', '水上人家'],
        ['槟城', '升旗山（Penang Hill）', 'RM30', '缆车上下，俯瞰槟城'],
        ['槟城', '极乐寺', '免费', '东南亚最大佛寺'],
        ['槟城', '新关仔角夜市', '免费', '槟城最著名夜市'],
        ['兰卡威', '天空之桥（Sky Bridge）', 'RM85含缆车', '高空玻璃桥'],
        ['兰卡威', '珍南海滩', '免费', '最热闹海滩'],
        ['兰卡威', '跳岛游', 'RM100-150', '孕妇岛、喂鹰'],
        ['兰卡威', '巨鹰广场', '免费', '兰卡威地标']
    ]
    add_table_from_data(doc, my_attr_headers, my_attr_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '3.2 餐厅推荐', level=2)
    doc.add_paragraph()
    
    my_rest_headers = ['城市', '餐厅名称', '人均', '特色/推荐菜']
    my_rest_data = [
        ['吉隆坡', 'Jalan Alor黄亚华', 'RM30-50', '烧鸡翅、福建面'],
        ['吉隆坡', 'Lot 10 Hutong', 'RM20-40', '美食广场，十号胡同'],
        ['吉隆坡', 'Madam Kwan\'s', 'RM30-50', '椰浆饭'],
        ['槟城', '伍秀泽海南鸡饭', 'RM15-25', ''],
        ['槟城', '潮州煎蕊（槟榔律）', 'RM5-10', ''],
        ['槟城', '德成饭店', 'RM40-60', '']
    ]
    add_table_from_data(doc, my_rest_headers, my_rest_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '3.3 酒店推荐', level=2)
    doc.add_paragraph()
    
    my_hotel_headers = ['城市', '酒店名称', '价位', '特色/备注']
    my_hotel_data = [
        ['吉隆坡', 'The Face Suites', '500-800元/晚', '无边泳池看双子塔'],
        ['吉隆坡', 'Traders Hotel', '800-1200元/晚', '双子塔景观房'],
        ['槟城', 'Eastern & Oriental Hotel', '600-1000元/晚', '殖民风格'],
        ['兰卡威', 'The Datai', '3000-6000元/晚', '顶级度假村'],
        ['兰卡威', 'Pelangi Beach Resort', '800-1500元/晚', '家庭友好']
    ]
    add_table_from_data(doc, my_hotel_headers, my_hotel_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '3.4 交通价格', level=2)
    doc.add_paragraph()
    
    my_trans_headers = ['交通方式', '价格', '备注']
    my_trans_data = [
        ['吉隆坡机场快线（KLIA Ekspres）', 'RM55', '28分钟到市中心'],
        ['Grab', '', '主要出行方式，便宜方便'],
        ['吉隆坡-槟城机票', 'RM100-300', '提前订便宜'],
        ['吉隆坡-槟城大巴', 'RM35-50', '约5小时'],
        ['兰卡威租车', 'RM80-150/天', '右舵驾驶']
    ]
    add_table_from_data(doc, my_trans_headers, my_trans_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '3.5 避坑清单', level=2)
    doc.add_paragraph()
    
    my_scam_headers = ['坑点', '对策']
    my_scam_data = [
        ['出租车不打表', '用Grab'],
        ['机场换汇汇率差', '少换一点，去市中心换'],
        ['购物退税门槛', '满RM300可退，同一店铺同一天'],
        ['榴莲不能带进酒店', '会被罚款']
    ]
    add_table_from_data(doc, my_scam_headers, my_scam_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '3.6 购物清单', level=2)
    doc.add_paragraph()
    
    my_shop_headers = ['类别', '推荐商品', '推荐购买地点']
    my_shop_data = [
        ['兰卡威免税', '巧克力（费列罗、好时）、烟酒', '兰卡威全岛免税'],
        ['白咖啡', 'Old Town、怡保白咖啡', '超市'],
        ['肉骨茶包', 'A1、奇香', '超市'],
        ['咖喱叻沙面', '槟城特产', '超市'],
        ['锡器', 'Royal Selangor', '专卖店'],
        ['兰花香水', 'Legendary香水', '专卖店']
    ]
    add_table_from_data(doc, my_shop_headers, my_shop_data)
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ==================== 新加坡 ====================
    add_heading_custom(doc, '四、新加坡攻略', level=1)
    
    add_heading_custom(doc, '4.1 核心景点', level=2)
    doc.add_paragraph()
    
    sg_attr_headers = ['景点名称', '门票/价格', '特色/备注']
    sg_attr_data = [
        ['滨海湾花园', '免费外观，云雾林+花穹SGD32', '灯光秀19:45/20:45'],
        ['鱼尾狮公园', '免费', '新加坡地标'],
        ['圣淘沙岛', '免费入岛', '环球影城SGD82、S.E.A.海洋馆SGD41'],
        ['克拉码头', '免费', '夜生活、酒吧'],
        ['牛车水（唐人街）', '免费', '美食、佛牙寺'],
        ['小印度', '免费', '印度风情、维拉玛卡里曼兴都庙'],
        ['乌节路', '免费', '购物天堂'],
        ['夜间动物园', 'SGD50', '世界首创'],
        ['摩天轮', 'SGD40', '亚洲最大'],
        ['樟宜机场', '免费', '星耀樟宜、室内瀑布']
    ]
    add_table_from_data(doc, sg_attr_headers, sg_attr_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '4.2 餐厅推荐', level=2)
    doc.add_paragraph()
    
    sg_rest_headers = ['餐厅名称', '人均', '特色/推荐菜', '备注']
    sg_rest_data = [
        ['了凡油鸡饭面（牛车水）', 'SGD5-10', '油鸡饭', '米其林一星'],
        ['天天海南鸡饭（Maxwell）', 'SGD5-8', '海南鸡饭', ''],
        ['松发肉骨茶', 'SGD10-15', '肉骨茶', '多分店'],
        ['珍宝海鲜（Jumbo）', 'SGD80-120', '辣椒蟹', ''],
        ['老巴刹（Lau Pa Sat）', 'SGD10-20', '熟食中心', ''],
        ['亚坤咖椰吐司', 'SGD5-8', '咖椰吐司', '早餐首选']
    ]
    add_table_from_data(doc, sg_rest_headers, sg_rest_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '4.3 酒店推荐', level=2)
    doc.add_paragraph()
    
    sg_hotel_headers = ['类型', '酒店名称', '价位', '特色/备注']
    sg_hotel_data = [
        ['经济型', 'Hotel 81系列', '400-600元/晚', '多分店'],
        ['中档型', 'V Hotel Lavender', '800-1200元/晚', '近地铁'],
        ['高档型', 'Marina Bay Sands', '3000-5000元/晚', '无边泳池'],
        ['奢华型', 'Raffles Hotel', '4000-8000元/晚', '百年传奇'],
        ['圣淘沙', 'Resorts World Sentosa', '2000-4000元/晚', '度假首选']
    ]
    add_table_from_data(doc, sg_hotel_headers, sg_hotel_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '4.4 交通价格', level=2)
    doc.add_paragraph()
    
    sg_trans_headers = ['交通方式', '价格', '备注']
    sg_trans_data = [
        ['地铁（MRT）', 'SGD0.9-2.2', '按距离计费'],
        ['巴士', 'SGD0.9-2.2', '与地铁通用'],
        ['出租车', 'SGD3.2起步', '每公里SGD0.55'],
        ['Grab', '比出租车便宜10-20%', ''],
        ['EZ-Link卡', 'SGD12（含SGD7余额）', '必备'],
        ['新加坡-马来西亚巴士', 'SGD20-40', '去新山']
    ]
    add_table_from_data(doc, sg_trans_headers, sg_trans_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '4.5 避坑清单', level=2)
    doc.add_paragraph()
    
    sg_scam_headers = ['坑点', '后果', '对策']
    sg_scam_data = [
        ['口香糖禁令', '禁止销售和进口', '违者罚款'],
        ['地铁禁食', '罚款SGD500', ''],
        ['乱丢垃圾', '罚款SGD300-1000', ''],
        ['吸烟区限制', '只能在指定区域吸烟', '违者罚款'],
        ['榴莲不能带上地铁/酒店', '会被罚款', ''],
        ['出租车附加费', '深夜、高峰、机场有附加费', '提前问清']
    ]
    add_table_from_data(doc, sg_scam_headers, sg_scam_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '4.6 购物清单', level=2)
    doc.add_paragraph()
    
    sg_shop_headers = ['类别', '推荐商品', '推荐购买地点']
    sg_shop_data = [
        ['本土品牌', '小CK（Charles & Keith）、Pazzion', '各商场'],
        ['茶叶', 'TWG Tea、Bacha Coffee', '专卖店'],
        ['食品', '肉骨茶包（松发、黄亚细）、咖椰酱（亚坤）', '超市'],
        ['免税化妆品', '各品牌', '樟宜机场，比市区便宜'],
        ['退税', '满SGD100可退', '机场自助办理']
    ]
    add_table_from_data(doc, sg_shop_headers, sg_shop_data)
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ==================== 印尼 ====================
    add_heading_custom(doc, '五、印尼攻略', level=1)
    
    add_heading_custom(doc, '5.1 核心目的地', level=2)
    doc.add_paragraph()
    
    id_attr_headers = ['城市/区域', '景点名称', '门票/价格', '特色/备注']
    id_attr_data = [
        ['巴厘岛乌布', '乌布皇宫/市场', '免费', '梯田、猴林、瑜伽、艺术村'],
        ['巴厘岛乌布', '德格拉朗梯田', '免费', '周杰伦《稻香》MV取景地'],
        ['巴厘岛乌布', '圣猴森林公园', '80,000盾', '猴子多，注意物品'],
        ['巴厘岛库塔', '库塔海滩', '免费', '冲浪、日落、夜生活'],
        ['巴厘岛水明漾', '水明漾海滩', '免费', '时尚海滩、餐厅酒吧'],
        ['巴厘岛乌鲁瓦图', '乌鲁瓦图寺庙', '50,000盾', '悬崖寺庙、日落'],
        ['巴厘岛金巴兰', '金巴兰海滩', '免费', '海滩烧烤日落'],
        ['巴厘岛努沙杜瓦', '努沙杜瓦海滩', '免费', '高端度假区、安静'],
        ['蓝梦岛', '恶魔眼泪', '免费', '一日游，海浪壮观'],
        ['佩尼达岛', '精灵坠崖', '免费', '一日游，网红打卡'],
        ['雅加达', '独立广场（Monas）', '免费', '地标，可登顶'],
        ['雅加达', '老城区（Kota Tua）', '免费', '殖民建筑、博物馆']
    ]
    add_table_from_data(doc, id_attr_headers, id_attr_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '5.2 餐厅推荐', level=2)
    doc.add_paragraph()
    
    id_rest_headers = ['城市/区域', '餐厅名称', '人均', '特色/推荐菜']
    id_rest_data = [
        ['巴厘岛乌布', 'Locavore', 'RP800,000+', '亚洲最佳餐厅，创意菜'],
        ['巴厘岛乌布', 'Warung Biah Biah', 'RP100,000-200,000', '印尼菜'],
        ['巴厘岛水明漾', 'Sisterfields', 'RP200,000-300,000', '网红brunch'],
        ['巴厘岛金巴兰', 'Menega Cafe', 'RP500,000+', '海鲜烧烤，日落晚餐'],
        ['巴厘岛', 'Bebek Bengil', 'RP150,000-250,000', '脏鸭餐'],
        ['雅加达', 'Plataran Menteng', 'RP500,000+', '印尼fine dining']
    ]
    add_table_from_data(doc, id_rest_headers, id_rest_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '5.3 酒店推荐', level=2)
    doc.add_paragraph()
    
    id_hotel_headers = ['区域', '酒店名称', '价位', '特色/备注']
    id_hotel_data = [
        ['乌布', 'Bisma Eight', '1000-1800元/晚', '设计酒店'],
        ['乌布', 'Kamandalu Ubud', '2000-4000元/晚', '梯田景观'],
        ['水明漾', 'The Legian', '1500-3000元/晚', '海滩front'],
        ['乌鲁瓦图', 'Alila Villas Uluwatu', '5000-10000元/晚', '顶级悬崖酒店'],
        ['努沙杜瓦', 'Ayodya Resort', '1200-2500元/晚', '宫殿风格'],
        ['蓝梦岛', 'Sandy Bay Beach Club', '400-800元/晚', '海景']
    ]
    add_table_from_data(doc, id_hotel_headers, id_hotel_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '5.4 交通价格', level=2)
    doc.add_paragraph()
    
    id_trans_headers = ['交通方式', '价格', '备注']
    id_trans_data = [
        ['巴厘岛机场-乌布', '300,000-400,000盾', '约1.5小时'],
        ['巴厘岛包车', '500,000-800,000盾/天', '10小时'],
        ['巴厘岛租摩托车', '80,000-150,000盾/天', ''],
        ['Grab/Gojek', '便宜', '部分地区禁止'],
        ['蓝梦岛船票', '150,000-300,000盾', '往返'],
        ['佩尼达岛一日游', '700,000-1,200,000盾', '含接送']
    ]
    add_table_from_data(doc, id_trans_headers, id_trans_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '5.5 避坑清单', level=2)
    doc.add_paragraph()
    
    id_scam_headers = ['坑点', '对策']
    id_scam_data = [
        ['机场出租车宰客', '提前订接机，或用Grab'],
        ['换汇黑市', '去正规换汇点，数清钱'],
        ['水上项目乱开价', '提前谈好价格，确认包含什么'],
        ['猴子抢东西', '乌布猴林，眼镜手机食物收好'],
        ['祭品别踩', '地上小花盒是祭品，别踩'],
        ['饮用水', '必须喝瓶装水，自来水不能喝'],
        ['小费', '不是强制，但服务好可以给10%']
    ]
    add_table_from_data(doc, id_scam_headers, id_scam_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '5.6 购物清单', level=2)
    doc.add_paragraph()
    
    id_shop_headers = ['类别', '推荐商品', '价格/备注']
    id_shop_data = [
        ['咖啡', '猫屎咖啡（Kopi Luwak）', 'RP500,000-1,000,000/100g'],
        ['手工皂', '巴厘岛手工皂', 'RP50,000-100,000，鸡蛋花香味'],
        ['精油/香薰', '精油、香薰', 'Ubud市场购买'],
        ['藤编包', '藤编包', 'Rp200,000-500,000'],
        ['银饰', '银饰', '乌布银器村'],
        ['沙龙', 'Sarong', 'Rp50,000-150,000，可当围巾']
    ]
    add_table_from_data(doc, id_shop_headers, id_shop_data)
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ==================== 日本 ====================
    add_heading_custom(doc, '六、日本攻略', level=1)
    
    add_heading_custom(doc, '6.1 核心城市与景点', level=2)
    doc.add_paragraph()
    
    jp_attr_headers = ['城市', '景点名称', '门票/价格', '特色/备注']
    jp_attr_data = [
        ['东京', '浅草寺', '免费', '雷门、仲见世商店街'],
        ['东京', '东京晴空塔', '¥2100', '350米观景台'],
        ['东京', '涩谷十字路口', '免费', '最繁忙路口'],
        ['东京', '明治神宫', '免费', '都市中的森林神社'],
        ['东京', '新宿', '免费', '歌舞伎町、思い出横丁'],
        ['东京', '秋叶原', '免费', '动漫电器圣地'],
        ['东京', 'teamLab Borderless', '¥3800', '数字艺术美术馆'],
        ['东京', '六本木之丘', '¥2200', '东京塔夜景'],
        ['京都', '清水寺', '¥400', '京都地标，清水舞台'],
        ['京都', '伏见稻荷大社', '免费', '千本鸟居'],
        ['京都', '金阁寺', '¥500', '金色舍利殿'],
        ['京都', '岚山', '免费', '竹林、渡月桥、小火车'],
        ['京都', '二年坂三年坂', '免费', '古街、抹茶'],
        ['京都', '琉璃光院', '¥2000', '秋季限定，绝美红叶'],
        ['大阪', '道顿堀', '免费', '格力高广告牌、美食'],
        ['大阪', '心斋桥', '免费', '购物天堂'],
        ['大阪', '大阪城', '¥600', '天守阁'],
        ['大阪', '环球影城', '¥8600起', '哈利波特、任天堂世界'],
        ['大阪', '黑门市场', '免费', '海鲜、和牛']
    ]
    add_table_from_data(doc, jp_attr_headers, jp_attr_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '6.2 餐厅推荐', level=2)
    doc.add_paragraph()
    
    jp_rest_headers = ['城市', '餐厅名称', '人均', '特色/推荐菜', '备注']
    jp_rest_data = [
        ['东京', '築地寿司大/大和', '¥3000-5000', '寿司', '排队1-2小时'],
        ['东京', '一兰拉面', '¥1200-1500', '拉面', '24小时'],
        ['东京', '鸟贵族', '¥2500-3500', '烧鸟', '连锁'],
        ['京都', '中村藤吉', '¥1500-2500', '抹茶甜品', ''],
        ['京都', '弘烧肉', '¥5000-8000', '和牛烤肉', ''],
        ['大阪', '蟹道乐', '¥8000-15000', '螃蟹料理', ''],
        ['大阪', '一兰/金龙拉面', '¥1000-1200', '拉面', ''],
        ['大阪', '黑门三平', '¥3000-5000', '海鲜丼', '']
    ]
    add_table_from_data(doc, jp_rest_headers, jp_rest_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '6.3 酒店推荐', level=2)
    doc.add_paragraph()
    
    jp_hotel_headers = ['城市', '酒店名称', '价位', '特色/备注']
    jp_hotel_data = [
        ['东京', 'Hotel Gracery Shinjuku', '800-1500元/晚', '哥斯拉酒店'],
        ['东京', 'Shibuya Excel Hotel Tokyu', '1000-1800元/晚', '涩谷站直连'],
        ['京都', 'Hotel Granvia Kyoto', '800-1500元/晚', '京都站直连'],
        ['京都', 'Gion Hatanaka', '1500-3000元/晚', '祇园附近，看艺伎'],
        ['大阪', 'Cross Hotel Osaka', '800-1500元/晚', '道顿堀中心'],
        ['大阪', 'Hotel Hankyu Respire', '600-1200元/晚', '梅田站直连']
    ]
    add_table_from_data(doc, jp_hotel_headers, jp_hotel_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '6.4 交通价格', level=2)
    doc.add_paragraph()
    
    jp_trans_headers = ['交通方式', '价格', '备注']
    jp_trans_data = [
        ['JR Pass（全国版7日）', '¥50,000', '新干线无限坐'],
        ['东京地铁24小时券', '¥800', '地铁无限坐'],
        ['大阪周游卡', '¥2800（1日）', '含景点+交通'],
        ['京都巴士一日券', '¥700', '巴士无限坐'],
        ['新干线东京-京都', '¥14,720', '约2小时15分'],
        ['新干线东京-大阪', '¥14,920', '约2小时30分'],
        ['西瓜卡（Suica）', '押金¥500', '必备交通卡'],
        ['出租车', '¥500起步', '每公里¥400，贵！']
    ]
    add_table_from_data(doc, jp_trans_headers, jp_trans_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '6.5 避坑清单', level=2)
    doc.add_paragraph()
    
    jp_scam_headers = ['坑点', '对策']
    jp_scam_data = [
        ['新干线自由席没座', '高峰期买指定席'],
        ['餐厅排队', '热门餐厅需预约或早到'],
        ['自动贩卖机陷阱', '看清楚再按，有些是按两次'],
        ['温泉礼仪', '必须先洗澡再泡，不能穿泳衣'],
        ['垃圾分类', '街上很少垃圾桶，随身携带垃圾袋'],
        ['吸烟区', '只能在指定区域吸烟'],
        ['小费', '日本不收小费，给反而失礼'],
        ['药妆店比价', '不同店价格差很多，多比较']
    ]
    add_table_from_data(doc, jp_scam_headers, jp_scam_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '6.6 购物清单', level=2)
    doc.add_paragraph()
    
    jp_shop_headers = ['类别', '推荐商品', '推荐购买地点']
    jp_shop_data = [
        ['药妆', '面膜、感冒药、眼药水、龙角散', '松本清、大国药妆'],
        ['零食', '白色恋人、Royce生巧、东京香蕉、薯条三兄弟', '机场、超市'],
        ['电器', '保温杯、吹风机、剃须刀', 'Bic Camera、友都八喜'],
        ['文具', 'MUJI、Loft、Itoya', ''],
        ['动漫周边', '动漫周边', '秋叶原、日本桥'],
        ['奢侈品', '各品牌', '银座、心斋桥，退税后便宜'],
        ['退税', '满¥5000可退', '现场办理']
    ]
    add_table_from_data(doc, jp_shop_headers, jp_shop_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '6.7 实用日语', level=2)
    doc.add_paragraph()
    
    jp_lang_headers = ['中文', '日语', '发音']
    jp_lang_data = [
        ['你好', 'Konnichiwa', '空尼奇瓦'],
        ['谢谢', 'Arigatou', '阿里嘎多'],
        ['对不起', 'Sumimasen', '斯米马森'],
        ['多少钱', 'Ikura desu ka', '衣库拉得斯卡'],
        ['好吃', 'Oishii', '哦伊西'],
        ['结账', 'Kaikei onegaishimasu', '开凯哦内嘎一西马斯'],
        ['1', 'ichi', '衣起'],
        ['2', 'ni', '尼'],
        ['3', 'san', '三']
    ]
    add_table_from_data(doc, jp_lang_headers, jp_lang_data)
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ==================== 韩国 ====================
    add_heading_custom(doc, '七、韩国攻略', level=1)
    
    add_heading_custom(doc, '7.1 核心城市与景点', level=2)
    doc.add_paragraph()
    
    kr_attr_headers = ['城市', '景点名称', '门票/价格', '特色/备注']
    kr_attr_data = [
        ['首尔', '景福宫', '₩3000', '穿韩服免费入场，周二闭馆'],
        ['首尔', '北村韩屋村', '免费', '传统韩屋，拍照圣地'],
        ['首尔', '明洞', '免费', '购物、美食'],
        ['首尔', '弘大', '免费', '年轻人聚集地，街头表演'],
        ['首尔', '梨泰院', '免费', '异国风情，夜生活'],
        ['首尔', 'N首尔塔', '₩11000', '南山夜景，爱情锁'],
        ['首尔', '乐天世界塔', '₩27000', '韩国最高建筑'],
        ['首尔', '广藏市场', '免费', 'Running Man取景地，生章鱼'],
        ['首尔', '东大门', '免费', '24小时批发市场'],
        ['釜山', '海云台海滩', '免费', '最著名海滩'],
        ['釜山', '甘川文化村', '免费', '彩色房子，小王子雕像'],
        ['釜山', '札嘎其市场', '免费', '海鲜市场，可现买现做'],
        ['釜山', 'BIFF广场', '免费', '电影节广场，小吃'],
        ['济州岛', '城山日出峰', '₩5000', '日出圣地，可爬山'],
        ['济州岛', '汉拿山', '免费', '韩国最高峰，需预约'],
        ['济州岛', '牛岛', '渡轮₩10500', '环岛骑行'],
        ['济州岛', '月汀里海滩', '免费', '白沙滩、咖啡馆'],
        ['济州岛', '泰迪熊博物馆', '₩12000', '']
    ]
    add_table_from_data(doc, kr_attr_headers, kr_attr_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '7.2 餐厅推荐', level=2)
    doc.add_paragraph()
    
    kr_rest_headers = ['城市', '餐厅名称', '人均', '特色/推荐菜']
    kr_rest_data = [
        ['首尔', '土俗村参鸡汤', '₩18000-25000', '景福宫附近'],
        ['首尔', '姜虎东白丁烤肉', '₩20000-35000', '连锁'],
        ['首尔', '王妃家烤肉', '₩30000-50000', '明洞'],
        ['首尔', 'BHC炸鸡', '₩20000-30000', '连锁'],
        ['釜山', '猪肉汤饭', '₩10000-15000', '札嘎其附近'],
        ['济州岛', '黑猪肉烤肉', '₩25000-40000', ''],
        ['济州岛', '海鲜锅', '₩30000-50000', '']
    ]
    add_table_from_data(doc, kr_rest_headers, kr_rest_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '7.3 酒店推荐', level=2)
    doc.add_paragraph()
    
    kr_hotel_headers = ['城市', '酒店名称', '价位', '特色/备注']
    kr_hotel_data = [
        ['首尔明洞', 'Loisir Hotel', '500-800元/晚', ''],
        ['首尔弘大', 'RYSE Autograph Collection', '1000-1800元/晚', '设计酒店'],
        ['首尔江南', 'Park Hyatt Seoul', '2000-4000元/晚', ''],
        ['釜山海云台', 'Park Hyatt Busan', '1500-3000元/晚', '海景'],
        ['济州岛', 'Jeju Shinhwa World', '1000-2000元/晚', '度假村'],
        ['济州岛', 'Hidden Cliff Hotel', '800-1500元/晚', '网红泳池']
    ]
    add_table_from_data(doc, kr_hotel_headers, kr_hotel_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '7.4 交通价格', level=2)
    doc.add_paragraph()
    
    kr_trans_headers = ['交通方式', '价格', '备注']
    kr_trans_data = [
        ['T-money卡', '₩3000押金', '地铁巴士通用'],
        ['首尔地铁', '₩1250起步', '按距离计费'],
        ['首尔-釜山KTX', '₩59800', '约2.5小时'],
        ['首尔-济州岛机票', '₩50000-150000', '约1小时'],
        ['济州岛包车', '₩150000-250000/天', ''],
        ['济州岛公交', '₩1200-3000/程', ''],
        ['出租车', '₩4800起步', '每公里₩100']
    ]
    add_table_from_data(doc, kr_trans_headers, kr_trans_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '7.5 避坑清单', level=2)
    doc.add_paragraph()
    
    kr_scam_headers = ['坑点', '对策']
    kr_scam_data = [
        ['出租车绕路', '用Kakao T，或看导航'],
        ['明洞购物宰客', '去免税店或明码标价店'],
        ['换汇', '明洞换汇所汇率好，多比较'],
        ['餐厅两人份起点', '很多烤肉店要求2人份起点'],
        ['垃圾分类', '街上垃圾桶少，带回酒店'],
        ['地铁老弱病残座', '粉色座位不要坐'],
        ['免税店提货', '机场提货，提前3小时到']
    ]
    add_table_from_data(doc, kr_scam_headers, kr_scam_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '7.6 购物清单', level=2)
    doc.add_paragraph()
    
    kr_shop_headers = ['类别', '推荐商品', '推荐购买地点']
    kr_shop_data = [
        ['化妆品', '面膜、彩妆、护肤品', 'Olive Young'],
        ['免税店', '各品牌', '乐天、新罗，提前办金卡'],
        ['零食', '蜂蜜黄油薯片、火鸡面、海苔', '超市'],
        ['服饰', 'ALAND、WONDER PLACE、8seconds', ''],
        ['文创', 'Line Friends、Kakao Friends', ''],
        ['人参', '正官庄', ''],
        ['退税', '满₩30000可退', '机场办理']
    ]
    add_table_from_data(doc, kr_shop_headers, kr_shop_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '7.7 实用韩语', level=2)
    doc.add_paragraph()
    
    kr_lang_headers = ['中文', '韩语', '发音']
    kr_lang_data = [
        ['你好', 'Annyeonghaseyo', '安宁哈塞哟'],
        ['谢谢', 'Gamsahamnida', '康桑哈密达'],
        ['对不起', 'Mianhamnida', '米安哈密达'],
        ['多少钱', 'Eolmaeyo', '哦儿马也哟'],
        ['好吃', 'Mashisseoyo', '马西搜哟'],
        ['结账', 'Kyesan haejuseyo', '凯三嗨租塞哟'],
        ['1', 'hana', '哈那'],
        ['2', 'dul', '读'],
        ['3', 'set', '塞']
    ]
    add_table_from_data(doc, kr_lang_headers, kr_lang_data)
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ==================== 通用工具 ====================
    add_heading_custom(doc, '八、通用工具', level=1)
    
    add_heading_custom(doc, '8.1 万能行李清单', level=2)
    doc.add_paragraph()
    
    checklist_headers = ['类别', '物品清单']
    checklist_data = [
        ['📋 证件类', '护照、签证/电子签、身份证、机票行程单、酒店预订单、保险单、驾照翻译件、2寸照片备用'],
        ['💳 财务类', '信用卡（Visa/Master）、银联卡、现金（当地货币+美元备用）、钱包、零钱包'],
        ['📱 电器类', '手机、充电器、充电宝（10000mAh以内可带上飞机）、转换插头（各国不同）、耳机、相机+存储卡'],
        ['👕 衣物类', '内衣裤（建议一次性）、袜子、T恤、长裤/长裙（寺庙用）、外套（空调房）、泳衣、拖鞋、舒适走路鞋'],
        ['🧴 洗漱类', '牙刷牙膏（部分酒店不提供）、洗发水、沐浴露、洗面奶、护肤品、防晒霜（必备）、剃须刀、梳子'],
        ['💊 药品类', '感冒药、止泻药、创可贴、晕车药、防蚊液、个人常用药'],
        ['🎒 其他', '雨伞/雨衣、墨镜、帽子、水杯、纸巾、湿巾、垃圾袋、笔记本+笔、U型枕、眼罩耳塞']
    ]
    add_table_from_data(doc, checklist_headers, checklist_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '8.2 紧急联系电话', level=2)
    doc.add_paragraph()
    
    emergency_headers = ['国家/机构', '电话']
    emergency_data = [
        ['中国外交部全球领保热线', '+86-10-12308'],
        ['泰国报警/旅游警察', '191 / 1155（旅游警察，有中文）'],
        ['泰国中国大使馆', '+66-2-245-7044'],
        ['越南报警', '113'],
        ['越南中国大使馆', '+84-24-3845-3736'],
        ['马来西亚报警/旅游警察', '999 / 03-2149-6590'],
        ['马来西亚中国大使馆', '+60-3-2164-5301'],
        ['新加坡报警', '999'],
        ['新加坡中国大使馆', '+65-6471-2117'],
        ['印尼报警', '110'],
        ['印尼中国大使馆', '+62-21-576-1039'],
        ['日本报警', '110'],
        ['日本中国大使馆', '+81-3-3403-3388'],
        ['韩国报警', '112'],
        ['韩国中国大使馆', '+82-2-738-1038'],
        ['Visa挂失', '+1-303-967-1096'],
        ['Mastercard挂失', '+1-636-722-7111'],
        ['银联挂失', '+86-21-6840-1888']
    ]
    add_table_from_data(doc, emergency_headers, emergency_data)
    doc.add_paragraph()
    
    add_heading_custom(doc, '8.3 汇率速查表', level=2)
    doc.add_paragraph()
    
    exchange_headers = ['货币', '汇率（约等于）']
    exchange_data = [
        ['泰铢（THB）', '1人民币 ≈ 5 泰铢'],
        ['越南盾（VND）', '1人民币 ≈ 3500 越南盾'],
        ['马来西亚林吉特（MYR）', '1人民币 ≈ 0.6 林吉特'],
        ['新加坡元（SGD）', '1人民币 ≈ 0.18 新元'],
        ['印尼盾（IDR）', '1人民币 ≈ 2200 印尼盾'],
        ['日元（JPY）', '1人民币 ≈ 20 日元'],
        ['韩元（KRW）', '1人民币 ≈ 190 韩元']
    ]
    add_table_from_data(doc, exchange_headers, exchange_data)
    doc.add_paragraph()
    
    # 保存文档
    doc.save('C:/Users/hua/Desktop/旅游攻略.docx')
    print("旅游攻略文档已生成：C:/Users/hua/Desktop/旅游攻略.docx")

if __name__ == '__main__':
    create_travel_guide()
