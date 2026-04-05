from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153) if level == 1 else RGBColor(0, 0, 0)
    return heading

def add_paragraph_custom(doc, text, bold=False, size=10.5, color=None):
    """添加自定义段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_bullet_list(doc, items):
    """添加项目符号列表"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.runs[0] if p.runs else p.add_run(item)
        if not p.runs:
            run.text = item
        run.font.name = '微软雅黑'
        run.font.size = Pt(10.5)

def add_numbered_list(doc, items):
    """添加编号列表"""
    for item in items:
        p = doc.add_paragraph(style='List Number')
        run = p.runs[0] if p.runs else p.add_run(item)
        if not p.runs:
            run.text = item
        run.font.name = '微软雅黑'
        run.font.size = Pt(10.5)

def create_travel_guide():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    
    # 封面标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('东南亚 · 日韩\n旅游攻略宝典')
    run.font.name = '微软雅黑'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('\n2025-2026 最新版\n一站式解决出行所有问题')
    run.font.name = '微软雅黑'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_page_break()
    
    # 目录
    add_heading_custom(doc, '目 录', level=1)
    toc_items = [
        '一、泰国旅游全攻略',
        '二、新加坡旅游全攻略', 
        '三、马来西亚旅游全攻略',
        '四、日本旅游全攻略',
        '五、韩国旅游全攻略',
        '六、实用工具与资源',
        '七、紧急联系方式'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # ==================== 泰国 ====================
    add_heading_custom(doc, '一、泰国旅游全攻略', level=1)
    
    add_heading_custom(doc, '1.1 最佳旅行时间', level=2)
    add_paragraph_custom(doc, '泰国分为三个季节，选择合适的时机出行体验更佳：')
    
    # 泰国季节表格
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '季节'
    hdr_cells[1].text = '时间'
    hdr_cells[2].text = '特点'
    
    data = [
        ['凉季（最佳）', '11月-2月', '气温20-30℃，降雨少，最舒适'],
        ['热季', '3月-5月', '气温35-40℃，泼水节（4月）期间最热闹'],
        ['雨季', '6月-10月', '阵雨为主，酒店价格低，游客少']
    ]
    for i, row_data in enumerate(data, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10)
    
    add_heading_custom(doc, '1.2 必去目的地', level=2)
    
    destinations = [
        ('曼谷', '大皇宫、卧佛寺、考山路、Chatuchak周末市场、湄南河夜游'),
        ('清迈', '古城寺庙、周日夜市、大象保护营、素贴山、宁曼路'),
        ('普吉岛', '芭东海滩、皮皮岛跳岛、西蒙人妖秀、丛林飞跃'),
        ('苏梅岛', '查汶海滩、涛岛潜水、安通国家公园、渔人村'),
        ('芭提雅', '真理寺、四方水上市场、蒂芬妮秀、珊瑚岛')
    ]
    
    for city, spots in destinations:
        p = doc.add_paragraph()
        run = p.add_run(f'{city}：')
        run.bold = True
        run.font.name = '微软雅黑'
        run.font.size = Pt(11)
        p.add_run(spots).font.name = '微软雅黑'
    
    add_heading_custom(doc, '1.3 签证与入境', level=2)
    add_bullet_list(doc, [
        '免签政策：中国公民持普通护照可免签入境30天（2024年3月起永久生效）',
        '入境材料：护照（有效期6个月以上）、往返机票、酒店预订单、1万泰铢/人现金（抽查）',
        '入境通道：走"护照查验"通道，无需填写入境卡',
        '延期：如需停留超过30天，可在泰国移民局申请延期7天'
    ])
    
    add_heading_custom(doc, '1.4 交通指南', level=2)
    add_paragraph_custom(doc, '【国内交通】', bold=True)
    add_bullet_list(doc, [
        '曼谷-清迈：飞机1小时（约300-500元），夜班火车12小时（约200元）',
        '曼谷-普吉：飞机1.5小时（约400-800元），大巴12小时',
        '曼谷-芭提雅：大巴2小时（约30元），包车约1500泰铢',
        '清迈-普吉：飞机2小时（约500-900元）'
    ])
    
    add_paragraph_custom(doc, '【当地交通】', bold=True)
    add_bullet_list(doc, [
        'Grab：东南亚版滴滴，可绑定支付宝，价格透明',
        'Bolt：比Grab便宜20-30%，需现金支付',
        '双条车/嘟嘟车：议价乘车，提前谈好价格',
        'BTS/MRT：曼谷轨道交通，避免堵车首选'
    ])
    
    add_heading_custom(doc, '1.5 预算参考（5天4晚）', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '项目'
    hdr[1].text = '经济型'
    hdr[2].text = '舒适型'
    
    budget_data = [
        ['机票', '1500-2500元', '2500-4000元'],
        ['住宿（4晚）', '400-800元', '1200-2500元'],
        ['餐饮', '300-500元', '800-1500元'],
        ['交通+门票', '300-500元', '600-1000元'],
        ['总计', '2500-4300元', '5100-9000元']
    ]
    for i, row_data in enumerate(budget_data, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10)
    
    add_heading_custom(doc, '1.6 实用Tips', level=2)
    add_bullet_list(doc, [
        '小费：酒店20泰铢/天，餐厅可找零不强制，按摩50-100泰铢',
        '电话卡：AIS信号最好，Happy卡性价比高，机场可买',
        '插座：两孔扁插，国内可直接使用，三孔需转换器',
        '货币：1人民币≈5泰铢，建议带现金在Super Rich换汇最划算',
        '禁忌：不要摸头、不要用脚指人、进寺庙需穿长裤长裙'
    ])
    
    doc.add_page_break()
    
    # ==================== 新加坡 ====================
    add_heading_custom(doc, '二、新加坡旅游全攻略', level=1)
    
    add_heading_custom(doc, '2.1 最佳旅行时间', level=2)
    add_paragraph_custom(doc, '新加坡全年炎热多雨，但6-8月相对干燥，是最佳旅行时间。避开12月-1月的雨季和春节高峰期。')
    
    add_heading_custom(doc, '2.2 必去景点', level=2)
    attractions = [
        ('滨海湾花园', '擎天大树灯光秀（19:45/20:45）、云雾林、花穹'),
        ('圣淘沙岛', '环球影城、S.E.A.海洋馆、时光之翼、西乐索海滩'),
        ('鱼尾狮公园', '新加坡地标，必拍打卡照'),
        ('牛车水', '唐人街，美食聚集地，佛牙寺龙华院'),
        ('小印度', '印度文化中心，色彩斑斓的建筑和香料市场'),
        ('克拉码头', '夜生活聚集地，河畔餐厅酒吧'),
        ('乌节路', '购物天堂，ION Orchard、高岛屋'),
        ('新加坡动物园', '世界最佳动物园之一，夜间野生动物园')
    ]
    for name, desc in attractions:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}：')
        run.bold = True
        run.font.name = '微软雅黑'
        p.add_run(desc).font.name = '微软雅黑'
    
    add_heading_custom(doc, '2.3 签证与入境', level=2)
    add_bullet_list(doc, [
        '签证：需提前办理，电子签3-5个工作日出签，费用约300元',
        '入境卡：出发前3天内在线填写SG Arrival Card（免费）',
        '入境通道：自助通关（e-Gate）或人工通道',
        '停留期：通常30天，具体以入境章为准'
    ])
    
    add_heading_custom(doc, '2.4 交通指南', level=2)
    add_bullet_list(doc, [
        '地铁MRT：最便捷，可购买EZ-Link卡或Tourist Pass（1日/3日通票）',
        '巴士：覆盖地铁不到的地方，可用EZ-Link卡',
        'Grab：打车软件，价格较贵但方便',
        '步行：市区景点集中，很多地方步行可达'
    ])
    
    add_heading_custom(doc, '2.5 预算参考（4天3晚）', level=2)
    add_paragraph_custom(doc, '新加坡消费较高，建议预算：')
    add_bullet_list(doc, [
        '经济型：4000-6000元（住青旅/经济酒店，吃食阁）',
        '舒适型：8000-12000元（住四星酒店，正常餐厅）',
        '奢华型：15000元以上（住滨海湾金沙，米其林餐厅）'
    ])
    
    add_heading_custom(doc, '2.6 省钱攻略', level=2)
    add_bullet_list(doc, [
        '吃食阁（Hawker Center）：海南鸡饭、肉骨茶、叻沙，5-10新币搞定一餐',
        '免费景点：鱼尾狮公园、滨海湾花园（户外部分）、植物园',
        '景点套票：购买Singapore Tourist Pass或Klook套票更划算',
        '购物退税：满100新币可退7%消费税（GST）'
    ])
    
    doc.add_page_break()
    
    # ==================== 马来西亚 ====================
    add_heading_custom(doc, '三、马来西亚旅游全攻略', level=1)
    
    add_heading_custom(doc, '3.1 最佳旅行时间', level=2)
    add_paragraph_custom(doc, '马来西亚分东马（沙巴、砂拉越）和西马（吉隆坡、槟城、兰卡威）。东马3-10月较好，西马全年可去，避开11月-1月雨季。')
    
    add_heading_custom(doc, '3.2 必去目的地', level=2)
    destinations_my = [
        ('吉隆坡', '双子塔、独立广场、茨厂街、阿罗街夜市、黑风洞'),
        ('槟城', '乔治市壁画街、升旗山、极乐寺、姓氏桥、槟城美食'),
        ('兰卡威', '天空之桥、珍南海滩、红树林游船、免税购物'),
        ('沙巴（亚庇）', '丹绒亚路日落、美人鱼岛、神山公园、红树林萤火虫'),
        ('仙本那', '卡帕莱水屋、马布岛、诗巴丹潜水、海上吉普赛人')
    ]
    for city, spots in destinations_my:
        p = doc.add_paragraph()
        run = p.add_run(f'{city}：')
        run.bold = True
        run.font.name = '微软雅黑'
        p.add_run(spots).font.name = '微软雅黑'
    
    add_heading_custom(doc, '3.3 签证与入境', level=2)
    add_bullet_list(doc, [
        '免签：中国公民免签30天（2023年12月起）',
        '入境材料：护照（有效期6个月以上）、往返机票、酒店订单、现金/信用卡',
        '入境通道：自助通关（e-Gate）或人工通道',
        '注意：沙巴和砂拉越有独立入境章，从西马飞东马需再次入境'
    ])
    
    add_heading_custom(doc, '3.4 预算参考', level=2)
    add_paragraph_custom(doc, '马来西亚性价比很高，是东南亚最划算的目的地之一：')
    add_bullet_list(doc, [
        '经济型：2000-3500元/周（住青旅/民宿，吃路边摊）',
        '舒适型：4000-6000元/周（住三星/四星酒店，正常餐厅）',
        '仙本那水屋：3000-8000元/人（2-3晚，含餐食）'
    ])
    
    add_heading_custom(doc, '3.5 美食推荐', level=2)
    add_bullet_list(doc, [
        '槟城：炒粿条、叻沙、福建面、煎蕊、豆蔻汁',
        '吉隆坡：肉骨茶、椰浆饭、印度煎饼、沙爹',
        '沙巴：生肉面、叻沙、海鲜、榴莲',
        '兰卡威：海鲜烧烤、马来咖喱'
    ])
    
    doc.add_page_break()
    
    # ==================== 日本 ====================
    add_heading_custom(doc, '四、日本旅游全攻略', level=1)
    
    add_heading_custom(doc, '4.1 最佳旅行时间', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '季节'
    hdr[1].text = '时间'
    hdr[2].text = '推荐目的地'
    
    jp_seasons = [
        ['樱花季', '3月下旬-4月上旬', '东京、京都、大阪、奈良'],
        ['黄金周', '4月底-5月初', '避开或提前预订'],
        ['夏季', '6-8月', '北海道、轻井泽、冲绳、花火大会'],
        ['红叶季', '11月', '京都、奈良、日光、箱根']
    ]
    for i, row_data in enumerate(jp_seasons, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10)
    
    add_heading_custom(doc, '4.2 经典路线推荐', level=2)
    
    add_paragraph_custom(doc, '【路线一：关东5日游】', bold=True)
    add_bullet_list(doc, [
        'Day1：东京（浅草寺、晴空塔、秋叶原）',
        'Day2：东京（涩谷、原宿、明治神宫、新宿）',
        'Day3：富士山一日游（河口湖、忍野八海、奥特莱斯）',
        'Day4：镰仓（大佛、江之电、高校前站、江之岛）',
        'Day5：东京购物（银座、表参道、机场）'
    ])
    
    add_paragraph_custom(doc, '【路线二：关西5日游】', bold=True)
    add_bullet_list(doc, [
        'Day1：大阪（道顿堀、心斋桥、大阪城）',
        'Day2：京都（清水寺、二年坂三年坂、伏见稻荷大社）',
        'Day3：京都（金阁寺、岚山、锦市场）',
        'Day4：奈良（东大寺、奈良公园、春日大社）',
        'Day5：大阪环球影城或购物'
    ])
    
    add_heading_custom(doc, '4.3 签证与入境', level=2)
    add_bullet_list(doc, [
        '单次签证：有效期3个月，停留15天，约300-400元',
        '三年多次：首次需去冲绳/东北三县，之后不限目的地',
        '五年多次：需较高收入证明',
        '入境：填写Visit Japan Web（提前在线申报）或纸质入境卡'
    ])
    
    add_heading_custom(doc, '4.4 交通攻略', level=2)
    add_paragraph_custom(doc, '【JR Pass（日本铁路通票）】', bold=True)
    add_bullet_list(doc, [
        '全国版7日：约1500元，适合跨城市长途旅行',
        '关东/关西地区版：约600-800元，适合区域内深度游',
        '购买：提前在国内购买兑换券，到日本后兑换',
        '注意：不适用于Nozomi（希望号）和Mizuho（瑞穗号）'
    ])
    
    add_paragraph_custom(doc, '【城市内交通】', bold=True)
    add_bullet_list(doc, [
        '东京：Suica/Pasmo卡，地铁+JR+巴士通用',
        '大阪：ICOCA卡，地铁网络发达',
        '京都：巴士一日券（700日元）最划算',
        '导航：Google Maps+换乘案内APP'
    ])
    
    add_heading_custom(doc, '4.5 预算参考（5天4晚）', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '项目'
    hdr[1].text = '经济型'
    hdr[2].text = '舒适型'
    
    jp_budget = [
        ['机票', '1500-2500元', '2500-4500元'],
        ['住宿（4晚）', '800-1500元', '2000-4000元'],
        ['交通（JR Pass等）', '600-1000元', '1000-1500元'],
        ['餐饮', '600-1000元', '1500-2500元'],
        ['总计', '3500-6000元', '7000-13000元']
    ]
    for i, row_data in enumerate(jp_budget, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10)
    
    add_heading_custom(doc, '4.6 实用Tips', level=2)
    add_bullet_list(doc, [
        '现金：日本仍大量使用现金，建议带5-10万日元',
        '退税：满5000日元可退税，商场直接办理',
        '插座：两孔扁插，电压100V，国内电器可直接使用',
        '网络：提前租WiFi或买流量卡',
        '礼仪：公共场所保持安静、排队、不边走边吃'
    ])
    
    doc.add_page_break()
    
    # ==================== 韩国 ====================
    add_heading_custom(doc, '五、韩国旅游全攻略', level=1)
    
    add_heading_custom(doc, '5.1 最佳旅行时间', level=2)
    add_paragraph_custom(doc, '韩国四季分明，春秋两季最舒适：')
    add_bullet_list(doc, [
        '春季（3-5月）：樱花盛开，汝矣岛、南山公园最美',
        '秋季（9-11月）：红叶季，内藏山、雪岳山赏枫',
        '夏季（6-8月）：炎热多雨，可去釜山海滩',
        '冬季（12-2月）：寒冷干燥，适合滑雪和购物'
    ])
    
    add_heading_custom(doc, '5.2 必去目的地', level=2)
    destinations_kr = [
        ('首尔', '景福宫、北村韩屋村、明洞、弘大、东大门、南山塔'),
        ('釜山', '海云台海滩、甘川文化村、札嘎其市场、太宗台'),
        ('济州岛', '汉拿山、城山日出峰、牛岛、柱状节理带、泰迪熊博物馆'),
        ('江原道', '春川南怡岛、束草、滑雪场（冬季）')
    ]
    for city, spots in destinations_kr:
        p = doc.add_paragraph()
        run = p.add_run(f'{city}：')
        run.bold = True
        run.font.name = '微软雅黑'
        p.add_run(spots).font.name = '微软雅黑'
    
    add_heading_custom(doc, '5.3 签证与入境', level=2)
    add_bullet_list(doc, [
        '免签：济州岛免签30天，首尔/釜山需签证',
        '签证类型：单次（90天有效，停留30天）、五年多次',
        '办理时间：5-7个工作日',
        '入境：填写入境卡，可自助通关（已登记指纹者）'
    ])
    
    add_heading_custom(doc, '5.4 交通指南', level=2)
    add_bullet_list(doc, [
        '首尔地铁：T-money卡，覆盖全市，有中文标识',
        'KTX高铁：首尔-釜山2.5小时，提前购票有折扣',
        '市内巴士：T-money卡可用，但地铁更方便',
        '济州岛：建议包车或租车（需国际驾照）'
    ])
    
    add_heading_custom(doc, '5.5 购物与退税', level=2)
    add_bullet_list(doc, [
        '免税店：乐天、新罗、现代，可网上下单机场取货',
        '明洞：化妆品、服饰集中地',
        '东大门：批发市场，晚上营业',
        '退税：满3万韩元可退，免税店直接免税',
        '支付宝/微信：大部分商店可用'
    ])
    
    add_heading_custom(doc, '5.6 预算参考（4天3晚）', level=2)
    add_bullet_list(doc, [
        '经济型：2500-4000元（住民宿/青旅，吃路边摊）',
        '舒适型：5000-8000元（住四星酒店，正常餐厅）',
        '购物型：8000元以上（含免税店购物）'
    ])
    
    doc.add_page_break()
    
    # ==================== 实用工具 ====================
    add_heading_custom(doc, '六、实用工具与资源', level=1)
    
    add_heading_custom(doc, '6.1 必备APP', level=2)
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'APP名称'
    hdr[1].text = '用途'
    hdr[2].text = '适用地区'
    
    apps = [
        ['Google Maps', '导航、查路线', '通用'],
        ['Grab', '打车', '东南亚'],
        ['Google翻译', '拍照翻译、语音翻译', '通用'],
        ['换乘案内', '日本交通查询', '日本'],
        ['Hyperdia', '日本列车时刻', '日本'],
        ['KakaoMap', '韩国导航', '韩国'],
        ['Booking/Agoda', '订酒店', '通用'],
        ['Klook', '买门票、一日游', '通用']
    ]
    for i, row_data in enumerate(apps, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10)
    
    add_heading_custom(doc, '6.2 汇率参考', level=2)
    add_bullet_list(doc, [
        '泰国：1人民币 ≈ 5泰铢（THB）',
        '新加坡：1人民币 ≈ 0.19新币（SGD）',
        '马来西亚：1人民币 ≈ 0.65马币（MYR）',
        '日本：1人民币 ≈ 20日元（JPY）',
        '韩国：1人民币 ≈ 185韩元（KRW）'
    ])
    
    add_heading_custom(doc, '6.3 行李清单', level=2)
    add_bullet_list(doc, [
        '证件：护照、签证、机票行程单、酒店预订单',
        '电器：手机、充电宝、转换插头（三孔国家需要）',
        '衣物：根据季节准备，空调房带薄外套',
        '药品：常用药、创可贴、肠胃药',
        '其他：防晒霜、雨伞、随身WiFi/流量卡'
    ])
    
    doc.add_page_break()
    
    # ==================== 紧急联系方式 ====================
    add_heading_custom(doc, '七、紧急联系方式', level=1)
    
    add_heading_custom(doc, '7.1 中国驻外使领馆', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '国家'
    hdr[1].text = '紧急电话'
    hdr[2].text = '备注'
    
    embassies = [
        ['泰国', '+66-2-245-7044', '24小时领事保护'],
        ['新加坡', '+65-6475-0165', '24小时领事保护'],
        ['马来西亚', '+60-3-2164-5301', '24小时领事保护'],
        ['日本', '+81-3-6450-2195', '24小时领事保护'],
        ['韩国', '+82-2-755-0572', '24小时领事保护']
    ]
    for i, row_data in enumerate(embassies, 1):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10)
    
    add_heading_custom(doc, '7.2 当地紧急电话', level=2)
    add_bullet_list(doc, [
        '通用：报警/急救/火警 112（东南亚部分国家）、119（日本韩国）、191（泰国）',
        '泰国：旅游警察 1155（中文服务）',
        '日本：报警 110，急救/火警 119',
        '韩国：报警 112，急救/火警 119'
    ])
    
    add_heading_custom(doc, '7.3 其他重要电话', level=2)
    add_bullet_list(doc, [
        '中国外交部全球领事保护热线：+86-10-12308',
        '信用卡挂失：Visa/Mastercard 全球客服',
        '保险公司：出行前购买的旅游险客服电话'
    ])
    
    # 结尾
    doc.add_paragraph()
    doc.add_paragraph()
    ending = doc.add_paragraph()
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ending.add_run('祝你旅途愉快！')
    run.font.name = '微软雅黑'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    return doc

if __name__ == '__main__':
    doc = create_travel_guide()
    output_path = r'C:/Users/hua/Desktop/旅游攻略.docx'
    doc.save(output_path)
    print(f'文档已保存到: {output_path}')
