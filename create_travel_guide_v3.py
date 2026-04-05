#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
东南亚+日韩旅游攻略生成器 v3
7个国家完整版：泰国、越南、马来西亚、新加坡、印尼、日本、韩国
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_obj = tcBorders.find(qn(f'w:{edge}'))
            if edge_obj is None:
                edge_obj = parse_xml(f'<w:{edge} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                tcBorders.append(edge_obj)
            edge_obj.set(qn('w:val'), kwargs[edge])

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 102, 204)
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 153, 76)
        run.font.bold = True
    else:
        run.font.size = Pt(12)
        run.font.bold = True
    return heading

def add_content_bold(doc, title, content):
    """添加加粗标题+内容"""
    p = doc.add_paragraph()
    run = p.add_run(f"{title}")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(f" {content}").font.size = Pt(11)

def create_travel_guide():
    doc = Document()
    
    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    
    # 封面
    title = doc.add_heading('东南亚 + 日韩旅游完全攻略', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 102, 204)
        run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('泰国 · 越南 · 马来西亚 · 新加坡 · 印尼 · 日本 · 韩国\n实用工具书 · 可直接查询使用')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    
    # 目录
    add_heading_custom(doc, '目录', level=1)
    toc_items = [
        '一、泰国攻略',
        '二、越南攻略', 
        '三、马来西亚攻略',
        '四、新加坡攻略',
        '五、印尼攻略',
        '六、日本攻略',
        '七、韩国攻略',
        '八、通用工具（行李清单、紧急联系）'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # ==================== 泰国 ====================
    add_heading_custom(doc, '一、泰国攻略', level=1)
    
    add_heading_custom(doc, '1.1 核心城市与景点', level=2)
    
    # 曼谷
    add_heading_custom(doc, '【曼谷】- 寺庙与现代交融', level=3)
    add_content_bold(doc, '必去景点：', '')
    attractions_bkk = [
        ('大皇宫+玉佛寺', '500泰铢，8:30-15:30，着装要求：长裤长裙，不能露肩'),
        ('卧佛寺', '200泰铢，8:00-18:30，曼谷最大寺庙，泰式按摩发源地'),
        ('郑王庙（黎明寺）', '100泰铢，8:00-18:00，网红拍照点，可俯瞰湄南河'),
        ('四面佛', '免费，6:00-22:00，位于Central World旁，香火最旺'),
        ('水门寺大佛', '免费，8:00-18:00，2022新建，69米高金色大佛'),
        ('ICONSIAM商场', '免费，10:00-22:00，室内水上市场，高端购物'),
        ('Jodd Fairs夜市', '免费，17:00-00:00，网红火山排骨、水果西施'),
        ('Chatuchak周末市场', '免费，周六日9:00-18:00，亚洲最大周末市场')
    ]
    for name, info in attractions_bkk:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    # 清迈
    add_heading_custom(doc, '【清迈】- 古城与慢生活', level=3)
    attractions_cnx = [
        ('塔佩门', '免费，古城东门，喂鸽子拍照，周日夜市起点'),
        ('双龙寺（素贴山）', '30泰铢，6:00-20:00，俯瞰清迈全城，金色佛塔'),
        ('宁曼路', '免费，文青聚集地，咖啡馆、买手店、网红餐厅'),
        ('周日夜市', '免费，周日16:00-22:00，古城主干道，手工艺品、小吃'),
        ('大象保护营', '1500-2500泰铢，半天/全天，给大象洗澡喂食，无表演'),
        ('夜间动物园', '800泰铢，17:00-22:00，世界最大夜间动物园'),
        ('丛林飞跃', '1500-3000泰铢，Skyline/Flight of Gibbon等，最长900米滑索'),
        ('拜县（需2天）', '免费，童话小镇，黄色小屋、草莓园、大树秋千')
    ]
    for name, info in attractions_cnx:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    # 普吉岛
    add_heading_custom(doc, '【普吉岛】- 海岛度假首选', level=3)
    attractions_hkt = [
        ('芭东海滩', '免费，最热闹海滩，酒吧街、江西冷商场'),
        ('卡伦/卡塔海滩', '免费，水质更好，适合冲浪、看日落'),
        ('皮皮岛一日游', '1200-2500泰铢，玛雅湾、猴子沙滩、浮潜'),
        ('皇帝岛/珊瑚岛', '1500-3000泰铢，水质清澈，深潜圣地'),
        ('西蒙人妖秀', '800泰铢，18:00/19:30/21:00，最正规人妖秀'),
        ('普吉老镇', '免费，彩色葡式建筑，米其林餐厅聚集'),
        ('神仙半岛', '免费，最佳日落观景点'),
        ('查龙寺', '免费，普吉最大寺庙，可放鞭炮祈福')
    ]
    for name, info in attractions_hkt:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    # 芭提雅
    add_heading_custom(doc, '【芭提雅】- 夜生活与海岛', level=3)
    attractions_pattaya = [
        ('蒂芬妮人妖秀', '1200泰铢，18:00/19:30/21:00，泰国最著名人妖秀'),
        ('格兰岛一日游', '500-1500泰铢，水上项目齐全，拖拽伞、香蕉船'),
        ('真理寺', '500泰铢，8:00-18:00，全木雕建筑，壮观'),
        ('四方水上市场', '200泰铢，9:00-20:00，《杜拉拉升职记》取景地'),
        ('风月步行街', '免费，18:00-02:00，酒吧、夜店聚集')
    ]
    for name, info in attractions_pattaya:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '1.2 餐厅推荐（分区域+价位）', level=2)
    
    # 曼谷餐厅
    add_heading_custom(doc, '【曼谷】', level=3)
    
    add_content_bold(doc, '💰 经济型（人均<200泰铢）：', '')
    budget_bkk = [
        ('Jay Fai（痣姐热炒）', '米其林街头美食，蟹肉煎蛋卷、冬阴功，人均150-250泰铢，需排队1-2小时'),
        ('Thip Samai（鬼门炒粉）', '最佳泰式炒河粉，人均100-150泰铢，16:00-02:00'),
        ('Krua Apsorn', '皇室认证泰餐，蟹肉煎蛋、冬阴功，人均200-300泰铢'),
        ('水门海南鸡饭（红/绿制服）', '网红海南鸡饭，人均60-80泰铢')
    ]
    for name, info in budget_bkk:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_content_bold(doc, '💰💰 中档型（人均500-1000泰铢）：', '')
    mid_bkk = [
        ('Nara Thai Cuisine', 'Central World 7楼，环境好，人均600-800泰铢'),
        ('Kub Kao Kub Pla', 'EmQuartier商场，创意泰餐，人均500-700泰铢'),
        ('Somboon Seafood（建兴酒家）', '咖喱蟹必点，人均800-1200泰铢，多分店'),
        ('The Local', '传统泰餐，环境优雅，人均800-1000泰铢')
    ]
    for name, info in mid_bkk:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_content_bold(doc, '💰💰💰 高档型（人均1500泰铢+）：', '')
    luxury_bkk = [
        ('Gaggan Anand', '亚洲最佳餐厅，创意印度菜，人均5000+泰铢，需预约'),
        ('Le Normandie', '文华东方酒店，法式料理，米其林二星，人均4000+泰铢'),
        ('Sirocco', '莲花酒店63楼，高空餐厅，人均3000+泰铢')
    ]
    for name, info in luxury_bkk:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    # 清迈餐厅
    add_heading_custom(doc, '【清迈】', level=3)
    cnx_restaurants = [
        ('Khao Soi Khun Yai', '最佳泰北咖喱面，人均50-80泰铢，仅午餐'),
        ('Tong Tem Toh', '泰北菜，烤猪肉、香肠，人均200-300泰铢'),
        ('Khao Soi Islam', '清真泰北咖喱面，人均80-120泰铢'),
        ('The House by Ginger', '创意泰餐，环境好，人均400-600泰铢'),
        ('Dash! Restaurant', '古城内，泰式+西式，人均300-500泰铢'),
        ('Cooking Love', '性价比高，人均200-350泰铢')
    ]
    for name, info in cnx_restaurants:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    # 普吉餐厅
    add_heading_custom(doc, '【普吉】', level=3)
    hkt_restaurants = [
        ('No.6 Restaurant', '芭东网红餐厅，泰餐，人均300-500泰铢，需排队'),
        ('Doo Dee Thai Food', '性价比高，人均200-350泰铢'),
        ('Kan Eang@Pier', '海鲜餐厅，查龙码头，人均800-1500泰铢'),
        ('Raya Restaurant', '普吉镇，传统泰餐，人均400-600泰铢'),
        ('One Chun Cafe', '普吉镇，米其林推荐，人均500-800泰铢')
    ]
    for name, info in hkt_restaurants:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '1.3 酒店推荐（分区域+价位）', level=2)
    
    # 曼谷酒店
    add_heading_custom(doc, '【曼谷】', level=3)
    
    add_content_bold(doc, '暹罗区（购物方便）', '')
    hotels_siam = [
        ('经济型：Lub d Bangkok Siam', '150-300元/晚，青旅风格，近BTS'),
        ('中档型：Holiday Inn Express Siam', '400-600元/晚，位置极佳'),
        ('高档型：Siam Kempinski Hotel', '1500-2500元/晚，直连商场，有泳池')
    ]
    for item in hotels_siam:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    add_content_bold(doc, '素坤逸区（夜生活丰富）', '')
    hotels_sukhumvit = [
        ('经济型：The Salil Hotel Sukhumvit 57', '300-500元/晚，网红酒店，拍照好看'),
        ('中档型：Grande Centre Point Sukhumvit 55', '600-900元/晚，公寓式，带厨房'),
        ('高档型：Marriott Marquis Queen\'s Park', '1000-1800元/晚，五星级，设施完善')
    ]
    for item in hotels_sukhumvit:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    add_content_bold(doc, '湄南河区（景观好）', '')
    hotels_river = [
        ('中档型：Chatrium Hotel Riverside', '500-800元/晚，河景房'),
        ('高档型：Mandarin Oriental Bangkok', '2500-5000元/晚，百年传奇酒店'),
        ('奢华型：The Peninsula Bangkok', '3000-6000元/晚，顶级服务')
    ]
    for item in hotels_river:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    # 清迈酒店
    add_heading_custom(doc, '【清迈】', level=3)
    hotels_cnx = [
        ('经济型：De Lanna Hotel', '150-250元/晚，古城内，泰式风格'),
        ('中档型：Buri Siri Boutique Hotel', '300-500元/晚，宁曼路附近，设计感强'),
        ('高档型：137 Pillars House', '1500-3000元/晚，殖民风格，服务顶级'),
        ('特色型：Chai Lai Orchid', '500-1000元/晚，大象民宿，与大象同住')
    ]
    for item in hotels_cnx:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    # 普吉酒店
    add_heading_custom(doc, '【普吉】', level=3)
    
    add_content_bold(doc, '芭东海滩（热闹便利）', '')
    hotels_patong = [
        ('经济型：Lub d Phuket Patong', '200-350元/晚，青旅+酒店'),
        ('中档型：Novotel Phuket Vintage Park', '500-800元/晚，近海滩'),
        ('高档型：Amari Phuket', '1200-2000元/晚，私人海滩')
    ]
    for item in hotels_patong:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    add_content_bold(doc, '卡伦/卡塔海滩（安静度假）', '')
    hotels_karon = [
        ('中档型：Centara Karon Resort', '400-700元/晚，家庭友好'),
        ('高档型：Sawasdee Village', '800-1500元/晚，泰式风格，泳池美'),
        ('奢华型：The Shore at Katathani', '2000-4000元/晚，仅限成人，私密性高')
    ]
    for item in hotels_karon:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '1.4 交通价格参考', level=2)
    
    transport_th = [
        ('曼谷BTS（轻轨）', '15-59泰铢/程，按距离计费'),
        ('曼谷MRT（地铁）', '16-42泰铢/程'),
        ('曼谷出租车', '35泰铢起步，每公里5-8泰铢，一定要打表（meter）'),
        ('Grab打车', '比出租车贵20-30%，但明码标价'),
        ('嘟嘟车（Tuk-tuk）', '100-300泰铢/程，需砍价'),
        ('摩的', '10-50泰铢/程，堵车时最快'),
        ('曼谷-清迈机票', '800-2000泰铢，提前订便宜'),
        ('曼谷-清迈大巴', '600-1000泰铢，VIP大巴约10小时'),
        ('曼谷-普吉机票', '1000-3000泰铢'),
        ('普吉机场-芭东海滩', '150-200泰铢（机场巴士），800-1000泰铢（出租车）'),
        ('清迈双条车', '20-50泰铢/人（红色），包车200-500泰铢'),
        ('清迈Grab', '比双条车贵，但方便'),
        ('普吉岛租摩托车', '200-400泰铢/天，需国际驾照或驾照翻译件')
    ]
    for item in transport_th:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item[0]}：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '1.5 避坑清单', level=2)
    
    scams_th = [
        ('大皇宫门口骗局', '有人说"大皇宫今天不开门"，骗你去其他景点消费。对策：无视，直接走到正门'),
        ('出租车不打表', '司机说"by meter broken"，要求一口价。对策：坚持打表，不打就换一辆'),
        ('嘟嘟车乱开价', '开口就要300-500泰铢。对策：提前查好距离，砍价到100-200'),
        ('水上市场/火车市场一日游陷阱', '报价极低（100泰铢），但会带去购物点。对策：选择正规旅行社，或自行前往'),
        ('人妖秀拍照收费', '表演结束后人妖主动拉你拍照，拍完要100泰铢小费。对策：不想拍就拒绝'),
        ('按摩店价格不透明', '街边按摩店可能临时加价。对策：选择连锁品牌如Let\'s Relax, Health Land'),
        ('租摩托车押金纠纷', '还车时以各种理由扣押金。对策：租车时拍照录像，确认划痕'),
        ('海鲜餐厅宰客', '芭东部分餐厅价格虚高。对策：看菜单明码标价，或去班赞海鲜市场买后加工')
    ]
    for name, info in scams_th:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '1.6 购物清单', level=2)
    
    shopping_th = [
        ('零食类', '小老板海苔、大哥花生、Bento鱿鱼片、MAMA泡面、泰式奶茶粉'),
        ('药品类', '青草膏、蜈蚣丸、虎牌膏药、鼻通、soffell驱蚊水'),
        ('化妆品', 'Mistine彩妆、Beauty Buffet牛奶系列、蜗牛霜'),
        ('服饰类', '曼谷包（NaRaYa）、泰丝、泰式服饰'),
        ('工艺品', '佛牌、木雕、银器、香薰'),
        ('推荐购物地', 'Boots（药品化妆品）、7-11（零食）、Big C（超市）、Chatuchak（手工艺品）、ICONSIAM（高端）'),
        ('退税', '满2000泰铢可退，VAT Refund标志的店，离境时在机场办理')
    ]
    for name, info in shopping_th:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '1.7 实用泰语', level=2)
    
    thai_phrases = [
        ('你好', 'Sa-wat-dee（萨瓦迪卡）'),
        ('谢谢', 'Khop-khun（考坤）'),
        ('对不起', 'Khor-thot（考托）'),
        ('多少钱', 'Tao-rai（套来）'),
        ('太贵了', 'Phaeng-mai（潘卖）'),
        ('便宜点', 'Lot-rai-hoi-noi（落来诺依）'),
        ('好吃', 'A-roy（阿洛依）'),
        ('厕所在哪', 'Hong-nam-yu-tee-nai（洪南由提乃）'),
        ('救命', 'Chuay-duay（求端）'),
        ('1-10数字', 'nung(1)、song(2)、sam(3)、see(4)、ha(5)、hok(6)、jet(7)、pad(8)、kao(9)、sip(10)')
    ]
    for thai, pron in thai_phrases:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{thai}：').bold = True
        p.add_run(pron)
    
    doc.add_page_break()
    
    # ==================== 越南 ====================
    add_heading_custom(doc, '二、越南攻略', level=1)
    
    add_heading_custom(doc, '2.1 核心城市与景点', level=2)
    
    add_heading_custom(doc, '【河内】- 法式殖民与越南风情', level=3)
    attractions_hanoi = [
        ('还剑湖', '免费，市中心地标，早晨有当地人晨练'),
        ('老城区（36行街）', '免费，每条街卖一种商品，摩托车穿行'),
        ('圣约瑟夫大教堂', '免费，河内版巴黎圣母院，哥特式建筑'),
        ('胡志明纪念堂', '免费，周二至周四、周六日8:00-11:00开放'),
        ('水上木偶戏', '100,000-200,000越南盾，Thang Long Water Puppet Theatre'),
        ('火车街', '免费，火车从居民楼间穿行，网红打卡点')
    ]
    for name, info in attractions_hanoi:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '【岘港】- 海滩与巴拿山', level=3)
    attractions_danang = [
        ('美溪海滩', '免费，世界最美六大海滩之一，白沙细腻'),
        ('巴拿山（Ba Na Hills）', '750,000越南盾，法式山城，佛手桥、游乐园'),
        ('会安古城（近岘港）', '免费入城，灯笼夜景绝美， tailor定制奥黛'),
        ('山茶半岛', '免费，观音像、看海、看猴'),
        ('五行山', '40,000越南盾，溶洞、佛教圣地')
    ]
    for name, info in attractions_danang:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '【芽庄】- 潜水与海岛', level=3)
    attractions_nhatrang = [
        ('四岛游', '200,000-400,000越南盾，浮潜、水上项目'),
        ('泥浆浴', '150,000-300,000越南盾，I-Resort或Thap Ba'),
        ('婆那加占婆塔', '22,000越南盾，印度教遗址，小吴哥窟'),
        ('钟屿石岬角', '22,000越南盾，法国电影《情人》取景地'),
        ('珍珠岛乐园', '880,000越南盾，跨海缆车、水上乐园')
    ]
    for name, info in attractions_nhatrang:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '【胡志明市】- 南部经济中心', level=3)
    attractions_sgn = [
        ('中央邮局', '免费，法式建筑，可寄明信片'),
        ('红教堂（圣母大教堂）', '免费，红砖外墙，对面就是邮局'),
        ('统一宫', '40,000越南盾，历史建筑，了解越战'),
        ('战争遗迹博物馆', '40,000越南盾，越战历史，震撼'),
        ('范五老街', '免费，背包客聚集地，酒吧街'),
        ('咖啡公寓', '免费外观，9层楼的咖啡馆集合体')
    ]
    for name, info in attractions_sgn:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '2.2 餐厅推荐', level=2)
    
    add_heading_custom(doc, '【河内】', level=3)
    restaurants_hanoi = [
        ('Pho Gia Truyen（小燕糯米饭）', '糯米饭配肉，人均30,000-50,000盾'),
        ('Bun Cha Huong Lien（奥巴马米粉）', '烤肉米粉，奥巴马同款，人均50,000-80,000盾'),
        ('Pho Thin', '牛肉河粉，人均40,000-60,000盾'),
        ('Cafe Giang', '鸡蛋咖啡发源地，人均25,000盾')
    ]
    for name, info in restaurants_hanoi:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '【岘港/会安】', level=3)
    restaurants_danang = [
        ('Mi Quang 1A', '广南面，岘港特色，人均40,000-60,000盾'),
        ('Banh Xeo Ba Duong', '越南煎饼，人均30,000-50,000盾'),
        ('Morning Glory（会安）', '高档越南菜，人均200,000-400,000盾'),
        ('White Rose Restaurant（会安）', '白玫瑰、炸云吞，人均50,000-100,000盾')
    ]
    for name, info in restaurants_danang:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '【芽庄】', level=3)
    restaurants_nhatrang = [
        ('Lac Canh Restaurant', '烧烤海鲜，人均150,000-300,000盾'),
        ('Nem Nuong Nha Trang', '芽庄春卷，人均50,000-100,000盾'),
        ('Yen\'s Restaurant', '越南菜，人均100,000-200,000盾')
    ]
    for name, info in restaurants_nhatrang:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '2.3 酒店推荐', level=2)
    
    hotels_vn = [
        ('【河内】Hanoi La Siesta Hotel', '400-700元/晚，老城区，服务超好'),
        ('【河内】Hilton Hanoi Opera', '800-1200元/晚，法式风格'),
        ('【岘港】Novotel Danang Premier Han River', '500-800元/晚，江景'),
        ('【岘港】InterContinental Danang Sun Peninsula Resort', '2000-4000元/晚，顶级度假'),
        ('【芽庄】Liberty Central Nha Trang', '300-500元/晚，海景'),
        ('【芽庄】Vinpearl Resort & Spa', '1000-2000元/晚，珍珠岛上'),
        ('【胡志明】Liberty Central Saigon Centre', '400-700元/晚，市中心'),
        ('【胡志明】Hotel des Arts Saigon', '1000-1800元/晚，法式奢华')
    ]
    for item in hotels_vn:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '2.4 交通价格', level=2)
    transport_vn = [
        ('河内/胡志明出租车', '12,000-15,000越南盾起步，每公里11,000-15,000盾'),
        ('Grab', '比出租车便宜，明码标价'),
        ('摩托车出租', '100,000-200,000越南盾/天'),
        ('Open Bus（长途巴士）', 'Futa、新咖啡，河内-胡志明约400,000-600,000盾'),
        ('越南国内航班', '胡志明-河内约800,000-1,500,000盾'),
        ('火车', '河内-岘港约400,000-800,000盾，卧铺')
    ]
    for item in transport_vn:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item[0]}：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '2.5 避坑清单', level=2)
    scams_vn = [
        ('出租车绕路/调表', '对策：用Grab，或提前查好路线'),
        ('找零少给', '越南盾面额大，容易数错。对策：当面数清'),
        ('摩托车抢包', '胡志明、河内常见。对策：包背内侧，手机别拿手上'),
        ('购物宰客', '没有明码标价的地方先问价。对策：去大商场或超市'),
        ('海关小费', '部分海关人员索要小费。对策：坚决不给，装听不懂'),
        ('假咖啡', '买咖啡豆注意辨别。对策：去Trung Nguyen等正规品牌店')
    ]
    for name, info in scams_vn:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '2.6 购物清单', level=2)
    shopping_vn = [
        ('咖啡', 'G7速溶、Trung Nguyen咖啡豆、猫屎咖啡'),
        ('零食', '椰子糖、榴莲饼、腰果、蔬果干'),
        ('工艺品', '奥黛、漆器、竹编、丝绸'),
        ('药品', '白虎膏、蛇油'),
        ('推荐购物地', 'Big C超市、Co.opmart、Ben Thanh市场（胡志明）、会安夜市')
    ]
    for name, info in shopping_vn:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '2.7 实用越南语', level=2)
    viet_phrases = [
        ('你好', 'Xin chao（新交）'),
        ('谢谢', 'Cam on（感恩）'),
        ('多少钱', 'Bao nhieu tien（包尼乌田）'),
        ('太贵了', 'Dat qua（达夸）'),
        ('便宜点', 'Re hon di（热昏迪）'),
        ('好吃', 'Ngon（嗯）'),
        ('1-5数字', 'mot(1)、hai(2)、ba(3)、bon(4)、nam(5)')
    ]
    for viet, pron in viet_phrases:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{viet}：').bold = True
        p.add_run(pron)
    
    doc.add_page_break()
    
    # 继续添加其他国家...
    # 由于篇幅限制，这里添加马来西亚、新加坡、印尼、日本、韩国的精简版
    
    # ==================== 马来西亚 ====================
    add_heading_custom(doc, '三、马来西亚攻略', level=1)
    
    add_heading_custom(doc, '3.1 核心城市与景点', level=2)
    
    add_heading_custom(doc, '【吉隆坡】- 双子塔与现代都市', level=3)
    attractions_kl = [
        ('双子塔（Petronas Towers）', '免费外观，登顶RM80，需预约'),
        ('独立广场', '免费，殖民建筑，拍照打卡'),
        ('茨厂街（唐人街）', '免费，美食、购物'),
        ('黑风洞（Batu Caves）', '免费，272级台阶，彩虹阶梯'),
        ('阿罗街（Jalan Alor）', '免费，夜市美食街'),
        ('国家清真寺', '免费，9:00-18:00，需穿长袍'),
        ('KLCC公园', '免费，双子塔前，音乐喷泉')
    ]
    for name, info in attractions_kl:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '【槟城】- 美食与壁画', level=3)
    attractions_penang = [
        ('乔治市壁画街', '免费，网红壁画打卡'),
        ('姓氏桥', '免费，水上人家'),
        ('升旗山（Penang Hill）', 'RM30，缆车上下，俯瞰槟城'),
        ('极乐寺', '免费，东南亚最大佛寺'),
        ('新关仔角夜市', '免费，槟城最著名夜市')
    ]
    for name, info in attractions_penang:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '【兰卡威】- 免税海岛', level=3)
    attractions_langkawi = [
        ('天空之桥（Sky Bridge）', 'RM85含缆车，高空玻璃桥'),
        ('珍南海滩', '免费，最热闹海滩'),
        ('跳岛游', 'RM100-150，孕妇岛、喂鹰'),
        ('巨鹰广场', '免费，兰卡威地标'),
        ('免税购物', '巧克力、烟酒超便宜')
    ]
    for name, info in attractions_langkawi:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '3.2 餐厅推荐', level=2)
    restaurants_my = [
        ('【吉隆坡】Jalan Alor黄亚华', '烧鸡翅、福建面，人均RM30-50'),
        ('【吉隆坡】Lot 10 Hutong', '美食广场，十号胡同，人均RM20-40'),
        ('【吉隆坡】Madam Kwan\'s', '椰浆饭，人均RM30-50'),
        ('【槟城】伍秀泽海南鸡饭', '人均RM15-25'),
        ('【槟城】潮州煎蕊（槟榔律）', '人均RM5-10'),
        ('【槟城】德成饭店', '人均RM40-60')
    ]
    for item in restaurants_my:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '3.3 酒店推荐', level=2)
    hotels_my = [
        ('【吉隆坡】The Face Suites', '500-800元/晚，无边泳池看双子塔'),
        ('【吉隆坡】Traders Hotel', '800-1200元/晚，双子塔景观房'),
        ('【槟城】Eastern & Oriental Hotel', '600-1000元/晚，殖民风格'),
        ('【兰卡威】The Datai', '3000-6000元/晚，顶级度假村'),
        ('【兰卡威】Pelangi Beach Resort', '800-1500元/晚，家庭友好')
    ]
    for item in hotels_my:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '3.4 交通价格', level=2)
    transport_my = [
        ('吉隆坡机场快线（KLIA Ekspres）', 'RM55，28分钟到市中心'),
        ('Grab', '主要出行方式，便宜方便'),
        ('吉隆坡-槟城机票', 'RM100-300，提前订便宜'),
        ('吉隆坡-槟城大巴', 'RM35-50，约5小时'),
        ('兰卡威租车', 'RM80-150/天，右舵驾驶')
    ]
    for item in transport_my:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item[0]}：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '3.5 避坑清单', level=2)
    scams_my = [
        ('出租车不打表', '对策：用Grab'),
        ('机场换汇汇率差', '对策：少换一点，去市中心换'),
        ('购物退税门槛', '满RM300可退，同一店铺同一天'),
        ('榴莲不能带进酒店', '会被罚款')
    ]
    for name, info in scams_my:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '3.6 购物清单', level=2)
    shopping_my = [
        ('兰卡威免税', '巧克力（费列罗、好时）、烟酒'),
        ('白咖啡', 'Old Town、怡保白咖啡'),
        ('肉骨茶包', 'A1、奇香'),
        ('咖喱叻沙面', '槟城特产'),
        ('锡器', 'Royal Selangor'),
        ('兰花香水', 'Legendary香水')
    ]
    for name, info in shopping_my:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    doc.add_page_break()
    
    # ==================== 新加坡 ====================
    add_heading_custom(doc, '四、新加坡攻略', level=1)
    
    add_heading_custom(doc, '4.1 核心景点', level=2)
    attractions_sg = [
        ('滨海湾花园', '免费外观，云雾林+花穹SGD32，灯光秀19:45/20:45'),
        ('鱼尾狮公园', '免费，新加坡地标'),
        ('圣淘沙岛', '免费入岛，环球影城SGD82、S.E.A.海洋馆SGD41'),
        ('克拉码头', '免费，夜生活、酒吧'),
        ('牛车水（唐人街）', '免费，美食、佛牙寺'),
        ('小印度', '免费，印度风情、维拉玛卡里曼兴都庙'),
        ('乌节路', '免费，购物天堂'),
        ('夜间动物园', 'SGD50，世界首创'),
        ('摩天轮', 'SGD40，亚洲最大'),
        ('樟宜机场', '免费，星耀樟宜、室内瀑布')
    ]
    for name, info in attractions_sg:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '4.2 餐厅推荐', level=2)
    restaurants_sg = [
        ('了凡油鸡饭面（牛车水）', '米其林一星，人均SGD5-10'),
        ('天天海南鸡饭（Maxwell）', '人均SGD5-8'),
        ('松发肉骨茶', '人均SGD10-15，多分店'),
        ('珍宝海鲜（Jumbo）', '辣椒蟹，人均SGD80-120'),
        ('老巴刹（Lau Pa Sat）', '熟食中心，人均SGD10-20'),
        ('亚坤咖椰吐司', '早餐首选，人均SGD5-8')
    ]
    for name, info in restaurants_sg:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '4.3 酒店推荐', level=2)
    hotels_sg = [
        ('经济型：Hotel 81系列', '400-600元/晚，多分店'),
        ('中档型：V Hotel Lavender', '800-1200元/晚，近地铁'),
        ('高档型：Marina Bay Sands', '3000-5000元/晚，无边泳池'),
        ('高档型：Raffles Hotel', '4000-8000元/晚，百年传奇'),
        ('圣淘沙：Resorts World Sentosa', '2000-4000元/晚，度假首选')
    ]
    for item in hotels_sg:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '4.4 交通价格', level=2)
    transport_sg = [
        ('地铁（MRT）', 'SGD0.9-2.2，按距离计费'),
        ('巴士', 'SGD0.9-2.2，与地铁通用'),
        ('出租车', 'SGD3.2起步，每公里SGD0.55'),
        ('Grab', '比出租车便宜10-20%'),
        ('EZ-Link卡', 'SGD12（含SGD7余额），必备'),
        ('新加坡-马来西亚巴士', 'SGD20-40，去新山')
    ]
    for item in transport_sg:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item[0]}：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '4.5 避坑清单', level=2)
    scams_sg = [
        ('口香糖禁令', '禁止销售和进口，违者罚款'),
        ('地铁禁食', '罚款SGD500'),
        ('乱丢垃圾', '罚款SGD300-1000'),
        ('吸烟区限制', '只能在指定区域吸烟，违者罚款'),
        ('榴莲不能带上地铁/酒店', '会被罚款'),
        ('出租车附加费', '深夜、高峰、机场有附加费，提前问清')
    ]
    for name, info in scams_sg:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '4.6 购物清单', level=2)
    shopping_sg = [
        ('小CK（Charles & Keith）', '新加坡本土品牌，比国内便宜'),
        ('Pazzion', '新加坡鞋牌'),
        ('TWG Tea', '高端茶叶'),
        ('Bacha Coffee', '摩洛哥风格咖啡'),
        ('肉骨茶包', '松发、黄亚细'),
        ('咖椰酱', '亚坤'),
        ('免税化妆品', '樟宜机场，比市区便宜'),
        ('退税', '满SGD100可退，机场自助办理')
    ]
    for name, info in shopping_sg:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    doc.add_page_break()
    
    # ==================== 印尼 ====================
    add_heading_custom(doc, '五、印尼攻略', level=1)
    
    add_heading_custom(doc, '5.1 核心目的地', level=2)
    
    add_heading_custom(doc, '【巴厘岛】- 最热门海岛', level=3)
    attractions_bali = [
        ('乌布（Ubud）', '梯田、猴林、瑜伽、艺术村'),
        ('库塔/水明漾（Kuta/Seminyak）', '冲浪、日落、夜生活'),
        ('乌鲁瓦图（Uluwatu）', '悬崖寺庙、日落、高端酒店'),
        ('金巴兰（Jimbaran）', '海滩烧烤日落'),
        ('努沙杜瓦（Nusa Dua）', '高端度假区、安静'),
        ('蓝梦岛（Nusa Lembongan）', '一日游，恶魔眼泪'),
        ('佩尼达岛（Nusa Penida）', '一日游，精灵坠崖、天仙裂痕'),
        ('天空之门（Lempuyang）', '网红拍照点，需早起'),
        ('巴图尔火山日出', '凌晨2点出发，日出+温泉'),
        ('德格拉朗梯田', '周杰伦《稻香》MV取景地')
    ]
    for name, info in attractions_bali:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '【雅加达】- 首都', level=3)
    attractions_jkt = [
        ('独立广场（Monas）', '地标，可登顶'),
        ('老城区（Kota Tua）', '殖民建筑、博物馆'),
        ('大清真寺（Istiqlal）', '东南亚最大清真寺'),
        ('Grand Indonesia商场', '高端购物')
    ]
    for name, info in attractions_jkt:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '5.2 餐厅推荐', level=2)
    restaurants_id = [
        ('【巴厘岛乌布】Locavore', '亚洲最佳餐厅，创意菜，人均RP800,000+'),
        ('【巴厘岛乌布】Warung Biah Biah', '印尼菜，人均RP100,000-200,000'),
        ('【巴厘岛水明漾】Sisterfields', '网红 brunch，人均RP200,000-300,000'),
        ('【巴厘岛金巴兰】Menega Cafe', '海鲜烧烤，日落晚餐，人均RP500,000+'),
        ('【巴厘岛】Bebek Bengil', '脏鸭餐，人均RP150,000-250,000'),
        ('【雅加达】Plataran Menteng', '印尼 fine dining，人均RP500,000+')
    ]
    for name, info in restaurants_id:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '5.3 酒店推荐', level=2)
    hotels_id = [
        ('【乌布】Bisma Eight', '1000-1800元/晚，设计酒店'),
        ('【乌布】Kamandalu Ubud', '2000-4000元/晚，梯田景观'),
        ('【水明漾】The Legian', '1500-3000元/晚，海滩front'),
        ('【乌鲁瓦图】Alila Villas Uluwatu', '5000-10000元/晚，顶级悬崖酒店'),
        ('【努沙杜瓦】Ayodya Resort', '1200-2500元/晚，宫殿风格'),
        ('【蓝梦岛】Sandy Bay Beach Club', '400-800元/晚，海景')
    ]
    for item in hotels_id:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '5.4 交通价格', level=2)
    transport_id = [
        ('巴厘岛机场-乌布', '300,000-400,000印尼盾，约1.5小时'),
        ('巴厘岛包车', '500,000-800,000印尼盾/天（10小时）'),
        ('巴厘岛租摩托车', '80,000-150,000印尼盾/天'),
        ('Grab/Gojek', '便宜，但部分地区禁止'),
        ('蓝梦岛船票', '150,000-300,000印尼盾往返'),
        ('佩尼达岛一日游', '700,000-1,200,000印尼盾含接送')
    ]
    for item in transport_id:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item[0]}：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '5.5 避坑清单', level=2)
    scams_id = [
        ('机场出租车宰客', '对策：提前订接机，或用Grab'),
        ('换汇黑市', '对策：去正规换汇点，数清钱'),
        ('水上项目乱开价', '对策：提前谈好价格，确认包含什么'),
        ('猴子抢东西', '乌布猴林，眼镜、手机、食物要收好'),
        ('祭品别踩', '地上小花盒是祭品，别踩'),
        ('饮用水', '必须喝瓶装水，自来水不能喝'),
        ('小费', '不是强制，但服务好可以给10%')
    ]
    for name, info in scams_id:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '5.6 购物清单', level=2)
    shopping_id = [
        ('猫屎咖啡', 'Kopi Luwak，约RP500,000-1,000,000/100g'),
        ('巴厘岛手工皂', '鸡蛋花香味，RP50,000-100,000'),
        ('精油/香薰', 'Ubud市场购买'),
        ('藤编包', 'Rp200,000-500,000'),
        ('银饰', '乌布银器村'),
        ('沙龙（Sarong）', 'Rp50,000-150,000，可当围巾')
    ]
    for name, info in shopping_id:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    doc.add_page_break()
    
    # ==================== 日本 ====================
    add_heading_custom(doc, '六、日本攻略', level=1)
    
    add_heading_custom(doc, '6.1 核心城市与景点', level=2)
    
    add_heading_custom(doc, '【东京】- 现代与传统', level=3)
    attractions_tokyo = [
        ('浅草寺', '免费，雷门、仲见世商店街'),
        ('东京晴空塔', '¥2100，350米观景台'),
        ('涩谷十字路口', '免费，最繁忙路口'),
        ('明治神宫', '免费，都市中的森林神社'),
        ('新宿', '免费，歌舞伎町、思い出横丁'),
        ('秋叶原', '免费，动漫电器圣地'),
        ('筑地场外市场', '免费，海鲜、寿司'),
        (' teamLab Borderless', '¥3800，数字艺术美术馆'),
        ('皇居', '免费，东御苑开放'),
        ('六本木之丘', '¥2200，东京塔夜景')
    ]
    for name, info in attractions_tokyo:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '【京都】- 千年古都', level=3)
    attractions_kyoto = [
        ('清水寺', '¥400，京都地标，清水舞台'),
        ('伏见稻荷大社', '免费，千本鸟居'),
        ('金阁寺', '¥500，金色舍利殿'),
        ('岚山', '免费，竹林、渡月桥、小火车'),
        ('二年坂三年坂', '免费，古街、抹茶'),
        ('八坂神社', '免费，祇园附近'),
        ('花见小路', '免费，艺伎可能出现'),
        ('锦市场', '免费，京都厨房'),
        ('琉璃光院', '¥2000，秋季限定，绝美红叶')
    ]
    for name, info in attractions_kyoto:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '【大阪】- 美食之都', level=3)
    attractions_osaka = [
        ('道顿堀', '免费，格力高广告牌、美食'),
        ('心斋桥', '免费，购物天堂'),
        ('大阪城', '¥600，天守阁'),
        ('环球影城', '¥8600起，哈利波特、任天堂世界'),
        ('黑门市场', '免费，海鲜、和牛'),
        ('通天阁', '¥900，新世界地标'),
        ('梅田蓝天大厦', '¥1500，空中庭园展望台')
    ]
    for name, info in attractions_osaka:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '6.2 餐厅推荐', level=2)
    restaurants_jp = [
        ('【东京】築地寿司大/大和', '排队1-2小时，人均¥3000-5000'),
        ('【东京】一兰拉面', '人均¥1200-1500，24小时'),
        ('【东京】鸟贵族', '烧鸟连锁，人均¥2500-3500'),
        ('【京都】中村藤吉', '抹茶甜品，人均¥1500-2500'),
        ('【京都】弘烧肉', '和牛烤肉，人均¥5000-8000'),
        ('【大阪】蟹道乐', '螃蟹料理，人均¥8000-15000'),
        ('【大阪】一兰/金龙拉面', '人均¥1000-1200'),
        ('【大阪】黑门三平', '海鲜丼，人均¥3000-5000')
    ]
    for name, info in restaurants_jp:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '6.3 酒店推荐', level=2)
    hotels_jp = [
        ('【东京】Hotel Gracery Shinjuku', '800-1500元/晚，哥斯拉酒店'),
        ('【东京】Shibuya Excel Hotel Tokyu', '1000-1800元/晚，涩谷站直连'),
        ('【京都】Hotel Granvia Kyoto', '800-1500元/晚，京都站直连'),
        ('【京都】Gion Hatanaka', '1500-3000元/晚，祇园附近，看艺伎'),
        ('【大阪】Cross Hotel Osaka', '800-1500元/晚，道顿堀中心'),
        ('【大阪】Hotel Hankyu Respire', '600-1200元/晚，梅田站直连')
    ]
    for item in hotels_jp:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '6.4 交通价格', level=2)
    transport_jp = [
        ('JR Pass（全国版7日）', '¥50,000，新干线无限坐'),
        ('东京地铁24小时券', '¥800，地铁无限坐'),
        ('大阪周游卡', '¥2800（1日），含景点+交通'),
        ('京都巴士一日券', '¥700，巴士无限坐'),
        ('新干线东京-京都', '¥14,720，约2小时15分'),
        ('新干线东京-大阪', '¥14,920，约2小时30分'),
        ('西瓜卡（Suica）', '押金¥500，必备交通卡'),
        ('出租车', '¥500起步，每公里¥400，贵！')
    ]
    for item in transport_jp:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item[0]}：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '6.5 避坑清单', level=2)
    scams_jp = [
        ('新干线自由席没座', '对策：高峰期买指定席'),
        ('餐厅排队', '热门餐厅需预约或早到'),
        ('自动贩卖机陷阱', '看清楚再按，有些是按两次'),
        ('温泉礼仪', '必须先洗澡再泡，不能穿泳衣'),
        ('垃圾分类', '街上很少垃圾桶，随身携带垃圾袋'),
        ('吸烟区', '只能在指定区域吸烟'),
        ('小费', '日本不收小费，给反而失礼'),
        ('药妆店比价', '不同店价格差很多，多比较')
    ]
    for name, info in scams_jp:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '6.6 购物清单', level=2)
    shopping_jp = [
        ('药妆', '面膜、感冒药、眼药水、龙角散'),
        ('零食', '白色恋人、Royce生巧、东京香蕉、薯条三兄弟'),
        ('电器', '保温杯、吹风机、剃须刀'),
        ('文具', 'MUJI、Loft、Itoya'),
        ('动漫周边', '秋叶原、日本桥'),
        ('奢侈品', '银座、心斋桥，退税后便宜'),
        ('退税', '满¥5000可退，现场办理')
    ]
    for name, info in shopping_jp:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '6.7 实用日语', level=2)
    japanese_phrases = [
        ('你好', 'Konnichiwa（空尼奇瓦）'),
        ('谢谢', 'Arigatou（阿里嘎多）'),
        ('对不起', 'Sumimasen（斯米马森）'),
        ('多少钱', 'Ikura desu ka（衣库拉得斯卡）'),
        ('好吃', 'Oishii（哦伊西）'),
        ('结账', 'Kaikei onegaishimasu（开凯哦内嘎一西马斯）'),
        ('1-5数字', 'ichi(1)、ni(2)、san(3)、yon/shi(4)、go(5)')
    ]
    for jp, pron in japanese_phrases:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{jp}：').bold = True
        p.add_run(pron)
    
    doc.add_page_break()
    
    # ==================== 韩国 ====================
    add_heading_custom(doc, '七、韩国攻略', level=1)
    
    add_heading_custom(doc, '7.1 核心城市与景点', level=2)
    
    add_heading_custom(doc, '【首尔】- 时尚与传统', level=3)
    attractions_seoul = [
        ('景福宫', '₩3000，穿韩服免费入场，周二闭馆'),
        ('北村韩屋村', '免费，传统韩屋，拍照圣地'),
        ('明洞', '免费，购物、美食'),
        ('弘大', '免费，年轻人聚集地，街头表演'),
        ('梨泰院', '免费，异国风情，夜生活'),
        ('江南', '免费，高端购物、整形'),
        ('N首尔塔', '₩11000，南山夜景，爱情锁'),
        ('乐天世界塔', '₩27000，韩国最高建筑'),
        ('广藏市场', '免费，Running Man取景地，生章鱼'),
        ('东大门', '免费，24小时批发市场'),
        ('三清洞', '免费，文艺咖啡馆'),
        ('仁寺洞', '免费，传统文化街')
    ]
    for name, info in attractions_seoul:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '【釜山】- 海滨城市', level=3)
    attractions_busan = [
        ('海云台海滩', '免费，最著名海滩'),
        ('甘川文化村', '免费，彩色房子，小王子雕像'),
        ('札嘎其市场', '免费，海鲜市场，可现买现做'),
        ('BIFF广场', '免费，电影节广场，小吃'),
        ('太宗台', '免费，悬崖海岸'),
        ('海东龙宫寺', '免费，海边寺庙')
    ]
    for name, info in attractions_busan:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '【济州岛】- 度假胜地', level=3)
    attractions_jeju = [
        ('城山日出峰', '₩5000，日出圣地，可爬山'),
        ('汉拿山', '免费，韩国最高峰，需预约'),
        ('牛岛', '渡轮₩10500，环岛骑行'),
        ('月汀里海滩', '免费，白沙滩、咖啡馆'),
        ('涉地可支', '免费，海岸风光'),
        ('泰迪熊博物馆', '₩12000'),
        ('柱状节理带', '₩2000，火山地貌')
    ]
    for name, info in attractions_jeju:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '7.2 餐厅推荐', level=2)
    restaurants_kr = [
        ('【首尔】土俗村参鸡汤', '景福宫附近，人均₩18000-25000'),
        ('【首尔】姜虎东白丁烤肉', '连锁，人均₩20000-35000'),
        ('【首尔】王妃家烤肉', '明洞，人均₩30000-50000'),
        ('【首尔】广藏市场生章鱼/绿豆饼', '人均₩10000-20000'),
        ('【首尔】BHC炸鸡', '连锁，人均₩20000-30000'),
        ('【釜山】猪肉汤饭', '札嘎其附近，人均₩10000-15000'),
        ('【济州岛】黑猪肉烤肉', '人均₩25000-40000'),
        ('【济州岛】海鲜锅', '人均₩30000-50000')
    ]
    for name, info in restaurants_kr:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '7.3 酒店推荐', level=2)
    hotels_kr = [
        ('【首尔明洞】Loisir Hotel', '500-800元/晚'),
        ('【首尔弘大】RYSE Autograph Collection', '1000-1800元/晚，设计酒店'),
        ('【首尔江南】Park Hyatt Seoul', '2000-4000元/晚'),
        ('【釜山海云台】Park Hyatt Busan', '1500-3000元/晚，海景'),
        ('【济州岛】Jeju Shinhwa World', '1000-2000元/晚，度假村'),
        ('【济州岛】Hidden Cliff Hotel', '800-1500元/晚，网红泳池')
    ]
    for item in hotels_kr:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item[0] + '：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '7.4 交通价格', level=2)
    transport_kr = [
        ('T-money卡', '₩3000押金，地铁巴士通用'),
        ('首尔地铁', '₩1250起步，按距离计费'),
        ('首尔-釜山KTX', '₩59800，约2.5小时'),
        ('首尔-济州岛机票', '₩50000-150000，约1小时'),
        ('济州岛包车', '₩150000-250000/天'),
        ('济州岛公交', '₩1200-3000/程'),
        ('出租车', '₩4800起步，每公里₩100')
    ]
    for item in transport_kr:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item[0]}：').bold = True
        p.add_run(item[1])
    
    add_heading_custom(doc, '7.5 避坑清单', level=2)
    scams_kr = [
        ('出租车绕路', '对策：用Kakao T，或看导航'),
        ('明洞购物宰客', '对策：去免税店或明码标价店'),
        ('换汇', '明洞换汇所汇率好，多比较'),
        ('餐厅两人份起点', '很多烤肉店要求2人份起点'),
        ('垃圾分类', '街上垃圾桶少，带回酒店'),
        ('地铁老弱病残座', '粉色座位不要坐'),
        ('免税店提货', '机场提货，提前3小时到')
    ]
    for name, info in scams_kr:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '7.6 购物清单', level=2)
    shopping_kr = [
        ('化妆品', 'Olive Young：面膜、彩妆、护肤品'),
        ('免税店', '乐天、新罗，提前办金卡'),
        ('零食', '蜂蜜黄油薯片、火鸡面、海苔'),
        ('服饰', 'ALAND、WONDER PLACE、8seconds'),
        ('文创', 'Line Friends、Kakao Friends'),
        ('人参/红参', '正官庄'),
        ('泡菜/海苔', '超市购买'),
        ('退税', '满₩30000可退，机场办理')
    ]
    for name, info in shopping_kr:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(info)
    
    add_heading_custom(doc, '7.7 实用韩语', level=2)
    korean_phrases = [
        ('你好', 'Annyeonghaseyo（安宁哈塞哟）'),
        ('谢谢', 'Gamsahamnida（康桑哈密达）'),
        ('对不起', 'Mianhamnida（米安哈密达）'),
        ('多少钱', 'Eolmaeyo（哦儿马也哟）'),
        ('好吃', 'Mashisseoyo（马西搜哟）'),
        ('结账', 'Kyesan haejuseyo（凯三嗨租塞哟）'),
        ('1-5数字', 'hana(1)、dul(2)、set(3)、net(4)、daseot(5)')
    ]
    for kr, pron in korean_phrases:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{kr}：').bold = True
        p.add_run(pron)
    
    doc.add_page_break()
    
    # ==================== 通用工具 ====================
    add_heading_custom(doc, '八、通用工具', level=1)
    
    add_heading_custom(doc, '8.1 万能行李清单', level=2)
    
    checklist_items = [
        ('📋 证件类', '护照、签证/电子签、身份证、机票行程单、酒店预订单、保险单、驾照翻译件、2寸照片备用'),
        ('💳 财务类', '信用卡（Visa/Master）、银联卡、现金（当地货币+美元备用）、钱包、零钱包'),
        ('📱 电器类', '手机、充电器、充电宝（10000mAh以内可带上飞机）、转换插头（各国不同）、耳机、相机+存储卡'),
        ('👕 衣物类', '内衣裤（建议一次性）、袜子、T恤、长裤/长裙（寺庙用）、外套（空调房）、泳衣、拖鞋、舒适走路鞋'),
        ('🧴 洗漱类', '牙刷牙膏（部分酒店不提供）、洗发水、沐浴露、洗面奶、护肤品、防晒霜（必备）、剃须刀、梳子'),
        ('💊 药品类', '感冒药、止泻药、创可贴、晕车药、防蚊液、个人常用药'),
        ('🎒 其他', '雨伞/雨衣、墨镜、帽子、水杯、纸巾、湿巾、垃圾袋、笔记本+笔、U型枕、眼罩耳塞')
    ]
    for category, items in checklist_items:
        p = doc.add_paragraph()
        p.add_run(category + '：').bold = True
        p.add_run(items)
    
    add_heading_custom(doc, '8.2 紧急联系', level=2)
    
    emergency_contacts = [
        ('中国外交部全球领保热线', '+86-10-12308'),
        ('泰国报警/旅游警察', '191 / 1155（旅游警察，有中文）'),
        ('泰国中国大使馆', '+66-2-245-7044'),
        ('越南报警', '113'),
        ('越南中国大使馆', '+84-24-3845-3736'),
        ('马来西亚报警/旅游警察', '999 / 03-2149-6590'),
        ('马来西亚中国大使馆', '+60-3-2164-5301'),
        ('新加坡报警', '999'),
        ('新加坡中国大使馆', '+65-6471-2117'),
        ('印尼报警', '110'),
        ('印尼中国大使馆', '+62-21-576-1039'),
        ('日本报警', '110'),
        ('日本中国大使馆', '+81-3-3403-3388'),
        ('韩国报警', '112'),
        ('韩国中国大使馆', '+82-2-738-1038'),
        ('信用卡挂失', 'Visa: +1-303-967-1096，Mastercard: +1-636-722-7111，银联: +86-21-6840-1888')
    ]
    for name, number in emergency_contacts:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}：').bold = True
        p.add_run(number)
    
    add_heading_custom(doc, '8.3 汇率速查（约等于）', level=2)
    
    exchange_rates = [
        ('1人民币 ≈ 5 泰铢（THB）', ''),
        ('1人民币 ≈ 3500 越南盾（VND）', ''),
        ('1人民币 ≈ 0.6 马来西亚林吉特（MYR）', ''),
        ('1人民币 ≈ 0.18 新加坡元（SGD）', ''),
        ('1人民币 ≈ 2200 印尼盾（IDR）', ''),
        ('1人民币 ≈ 20 日元（JPY）', ''),
        ('1人民币 ≈ 190 韩元（KRW）', '')
    ]
    for rate, _ in exchange_rates:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(rate)
    
    # 保存文档
    doc.save('C:/Users/hua/Desktop/旅游攻略.docx')
    print("旅游攻略文档已生成：C:/Users/hua/Desktop/旅游攻略.docx")

if __name__ == '__main__':
    create_travel_guide()
